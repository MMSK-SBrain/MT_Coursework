"""
Generate participant handbook as formatted .docx for printing.
Emotional Intelligence, Bias & Microaggression Awareness — Disha Initiative
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Colors
DARK_TEAL = RGBColor(0x00, 0x6D, 0x77)
MEDIUM_TEAL = RGBColor(0x00, 0x96, 0xA3)
LIGHT_TEAL_BG = "E0F4F4"
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CONFIDENTIAL_TEXT = "Confidential — For Facilitated Session Use Only | Reynlab Technologies Pvt Ltd"


def add_confidential_footer(doc):
    """Add confidential text as a paragraph at bottom of conceptual page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CONFIDENTIAL_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = LIGHT_GRAY
    run.font.italic = True


def add_page_break(doc):
    doc.add_page_break()


def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_TEAL
    return heading


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_TEAL
    run.font.italic = True
    return p


def add_body(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_blank_line(doc, label="", num_lines=1):
    for _ in range(num_lines):
        if label:
            p = doc.add_paragraph()
            run = p.add_run(f"{label} ")
            run.bold = True
            run.font.size = Pt(10)
            run2 = p.add_run("_" * 80)
            run2.font.color.rgb = LIGHT_GRAY
        else:
            p = doc.add_paragraph()
            run = p.add_run("_" * 90)
            run.font.color.rgb = LIGHT_GRAY


def add_answer_block(doc, question_label, lines=4):
    p = doc.add_paragraph()
    run = p.add_run(question_label)
    run.bold = True
    run.font.size = Pt(11)
    for _ in range(lines):
        p = doc.add_paragraph()
        run = p.add_run("_" * 90)
        run.font.color.rgb = LIGHT_GRAY


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

# Spacer
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Emotional Intelligence, Bias &\nMicroaggression Awareness")
run.font.size = Pt(28)
run.font.color.rgb = DARK_TEAL
run.bold = True

doc.add_paragraph()

add_subtitle(doc, "Disha Initiative — Building Inclusive Workspaces")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PARTICIPANT HANDBOOK")
run.font.size = Pt(16)
run.font.color.rgb = MEDIUM_GRAY
run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("7 April 2026")
run.font.size = Pt(14)
run.font.color.rgb = MEDIUM_GRAY

for _ in range(4):
    doc.add_paragraph()

# Confidentiality notice
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIALITY NOTICE")
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = MEDIUM_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This workbook is provided for use during facilitated sessions only. "
    "Contents are proprietary to Reynlab Technologies Pvt Ltd. "
    "Reproduction, distribution, or use outside the facilitated session "
    "is prohibited without written permission."
)
run.font.size = Pt(8)
run.font.color.rgb = LIGHT_GRAY
run.italic = True

# ═══════════════════════════════════════════════════════════════
# PAGE 1: WELCOME & AGENDA
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Welcome & Program Overview")

add_body(doc,
    "Welcome to the Disha Initiative. Today's program is designed to help you build "
    "practical skills in emotional intelligence, bias awareness, and inclusive communication. "
    "Disha is about empowering you with tools to understand yourself more deeply, recognise "
    "patterns that shape how we interact, and take meaningful action toward workspaces where "
    "everyone can thrive. We are glad you are here."
)

doc.add_paragraph()
add_title(doc, "Today's Agenda", level=2)

agenda_rows = [
    ["9:30 – 10:00", "Opening & EQ Prism Check-In"],
    ["10:00 – 10:50", "Module 1: Emotional Intelligence Essentials"],
    ["10:50 – 11:05", "Tea Break"],
    ["11:05 – 12:00", "Module 2: Understanding Bias"],
    ["12:00 – 12:45", "Module 3: Stereotypes"],
    ["12:45 – 1:30", "Lunch"],
    ["1:30 – 1:45", "Energizer & Bridge"],
    ["1:45 – 2:40", "Module 4: Recognizing Microaggressions"],
    ["2:40 – 2:55", "Tea Break"],
    ["2:55 – 3:50", "Module 5: Responding & Intervening"],
    ["3:50 – 4:40", "Module 6: From Awareness to Inclusive Action"],
    ["4:40 – 5:15", "Closing & EQ Prism Retake"],
]
add_table(doc, ["Time", "Session"], agenda_rows)

doc.add_paragraph()
add_title(doc, "Our Ground Rules", level=2)
for _ in range(5):
    add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 2: EQ PRISM CHECK-IN
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "My EQ Prism Check-In (Opening)")

add_body(doc,
    "The EQ Prism is a quick self-awareness tool that helps you notice your current "
    "stress-energy state using colour psychology. It takes just 60 seconds and gives you "
    "a snapshot of where you are right now."
)

doc.add_paragraph()
add_title(doc, "My Zone", level=2)

p = doc.add_paragraph()
run = p.add_run("Circle or write the zone that matches your result:")
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for zone in ["Grounded", "Balanced", "Activated", "Overloaded"]:
    run = p.add_run(f"    {zone}    ")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK_TEAL

doc.add_paragraph()
add_blank_line(doc, "My Zone:")

doc.add_paragraph()
add_title(doc, "My One Word", level=2)
add_blank_line(doc, "One word that describes how I am feeling right now:")

doc.add_paragraph()
add_title(doc, "Notes", level=2)
for _ in range(3):
    add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 3: MODULE 1 — GOLEMAN MODEL
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 1 — Emotional Intelligence Essentials")

add_body(doc,
    "Emotional intelligence is the ability to recognise, understand, manage, and effectively "
    "use emotions — in yourself and in your interactions with others."
)

add_title(doc, "Goleman's 4-Quadrant Model", level=2)

# Create a 2x2 table for the quadrants
quad_table = doc.add_table(rows=2, cols=2)
quad_table.alignment = WD_TABLE_ALIGNMENT.CENTER

quadrants = [
    ("SELF-AWARENESS", "Recognising your own emotions as they happen and understanding their impact on your thoughts and behaviour."),
    ("SELF-MANAGEMENT", "Managing your emotional reactions and impulses so you can respond thoughtfully rather than react impulsively."),
    ("SOCIAL AWARENESS", "Sensing and understanding the emotions of others, and reading the dynamics in a group or team."),
    ("RELATIONSHIP MANAGEMENT", "Using your awareness of your own and others' emotions to build strong connections and navigate conflict."),
]

for i, (title, desc) in enumerate(quadrants):
    row = i // 2
    col = i % 2
    cell = quad_table.rows[row].cells[col]
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEAL
    p2 = cell.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(9)
    set_cell_shading(cell, LIGHT_TEAL_BG)

doc.add_paragraph()

for quadrant_name in ["Self-Awareness", "Self-Management", "Social Awareness", "Relationship Management"]:
    add_answer_block(doc, f"{quadrant_name} — My personal example:", lines=2)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 4: MODULE 1 — VIDEO DISCUSSION
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 1 — Video Discussion")

for q in ["Q1:", "Q2:", "Q3:"]:
    add_answer_block(doc, q, lines=5)
    doc.add_paragraph()

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 5: MODULE 1 — TRIGGERS & STOP
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 1 — My Triggers at Work")

add_body(doc, "Think about the last month at work. What situations triggered strong emotions in you?")

trigger_rows = [
    ["1", "", "", ""],
    ["2", "", "", ""],
    ["3", "", "", ""],
    ["4", "", "", ""],
    ["5", "", "", ""],
]
add_table(doc, ["#", "Situation", "Emotion", "Intensity (1-5)"], trigger_rows)

doc.add_paragraph()
add_title(doc, "The STOP Technique", level=2)

stop_rows = [
    ["S", "Stop what you are doing."],
    ["T", "Take a breath."],
    ["O", "Observe your emotions."],
    ["P", "Proceed with awareness."],
]
add_table(doc, ["Letter", "Action"], stop_rows)

doc.add_paragraph()
add_blank_line(doc, "ONE regulation strategy I will try this week:")
add_blank_line(doc)
add_blank_line(doc, "When I will use it:")
add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 6: MODULE 2 — FOUR TYPES OF BIAS
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 2 — Understanding Bias")
add_title(doc, "Four Types of Bias", level=2)

biases = [
    ("1. Unconscious Bias", "Automatic associations and attitudes we hold about groups of people, formed through our experiences, upbringing, and exposure — without our conscious awareness."),
    ("2. Affinity Bias", "The tendency to favour people who are similar to us — in background, interests, appearance, or experience — because they feel familiar and comfortable."),
    ("3. Confirmation Bias", "The tendency to search for, interpret, and remember information that confirms what we already believe — while ignoring evidence that contradicts it."),
    ("4. Attribution Bias", 'The tendency to explain our own behaviour through circumstances ("I was tired") but explain others\' behaviour through character ("She\'s lazy").'),
]

for title, desc in biases:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_TEAL
    run.font.size = Pt(11)

    add_body(doc, desc, size=10)
    add_blank_line(doc, "My example:")
    doc.add_paragraph()

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 7: MODULE 2 — QUICK CHECK & VIDEO DISCUSSION
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 2 — Quick Check & Video Discussion")

add_title(doc, "Quick Check: Which bias is at play?", level=2)

for i in range(1, 5):
    add_answer_block(doc, f"Scenario {i}:", lines=2)

doc.add_paragraph()
add_title(doc, "Video Discussion", level=2)

for q in ["Q1:", "Q2:"]:
    add_answer_block(doc, q, lines=4)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 8: MODULE 2 — CASE STUDY #1
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 2 — Case Study Discussion #1")

for q in ["Q1:", "Q2:", "Q3:"]:
    add_answer_block(doc, q, lines=5)

doc.add_paragraph()
add_title(doc, "Bias Interruption Strategies", level=2)

strategies = [
    ("1. Slow down decisions.", "When you catch yourself making a quick judgement about someone, pause. Fast thinking is where bias lives."),
    ("2. Seek disconfirming evidence.", 'Deliberately look for information that challenges your initial impression. Ask: "What am I not seeing?"'),
    ("3. Use structured criteria.", "Evaluate people against clear, pre-defined criteria rather than gut feeling. Write your criteria before you evaluate."),
]

for title, desc in strategies:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(f" {desc}")
    run2.font.size = Pt(10)

doc.add_paragraph()
add_blank_line(doc, "Where I will apply this:")
add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 9: MODULE 3 — STEREOTYPES
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 3 — Stereotypes")

add_body(doc,
    "A stereotype is an oversimplified, fixed belief about a group of people. "
    "Stereotypes are not always hostile — sometimes the most limiting ones sound like compliments."
)

add_title(doc, "Stereotype Formation Cycle", level=2)

cycle_rows = [
    ["1. Exposure", "Encounter with a group or individual"],
    ["2. Generalisation", "One example = all members"],
    ["3. Reinforcement", "We notice what confirms the belief"],
    ['4. "Truth"', "It feels like a fact"],
]
add_table(doc, ["Stage", "What Happens"], cycle_rows)

doc.add_paragraph()
add_title(doc, "4 Sources of Stereotypes", level=2)

source_rows = [
    ["Family", 'Messages we absorbed growing up about "how things are"'],
    ["Media", "Representations and narratives that shape our expectations"],
    ["Workplace culture", "Norms about who belongs in which roles and at which levels"],
    ["Self (internalised)", 'Stereotypes we apply to our own identity — "People like me don\'t..."'],
]
add_table(doc, ["Source", "How It Works"], source_rows)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 10: MODULE 3 — VIDEO DISCUSSION
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 3 — Video Discussion")

for q in ["Q1:", "Q2:", "Q3:"]:
    add_answer_block(doc, q, lines=5)
    doc.add_paragraph()

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 11: MODULE 3 — CASE STUDY #2 & INHERITED SCRIPTS
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 3 — Case Study Discussion #2")

add_title(doc, "Discussion Notes", level=2)
for _ in range(6):
    add_blank_line(doc)

doc.add_paragraph()
add_title(doc, "My Inherited Scripts", level=2)
add_body(doc, 'What "scripts" have you absorbed about people, roles, or capabilities?')

scripts_rows = [
    ["1", "", "", ""],
    ["2", "", "", ""],
    ["3", "", "", ""],
]
add_table(doc, ["#", "Script", "Source (Family/Media/Work/Self)", "Still carrying it? (Y/N)"], scripts_rows)

doc.add_paragraph()
add_title(doc, "Connecting the Dots", level=2)
add_body(doc, "Trace one personal example through the chain: Bias → Stereotype → Emotion → Behaviour", italic=True)
for _ in range(4):
    add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 12: AFTERNOON BRIDGE
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Afternoon Bridge")

add_title(doc, "Morning Recap", level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EI (Understanding Self)  →  Bias (How We Think)  →  Stereotypes (What We Believe)")
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = DARK_TEAL

doc.add_paragraph()
add_body(doc,
    "This morning took you from understanding your own emotions, to recognising the invisible "
    "filters that shape your thinking, to examining the inherited beliefs that influence how you see others."
)

doc.add_paragraph()
add_blank_line(doc, "One thing that stuck with me from this morning:")
for _ in range(3):
    add_blank_line(doc)

doc.add_paragraph()
add_title(doc, "What's Ahead", level=2)
add_body(doc,
    "This afternoon, we move from awareness to action. You will learn to spot microaggressions, "
    "practise responding, and build your personal plan for inclusive behaviour."
)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 13: MODULE 4 — MICROAGGRESSIONS
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 4 — Recognizing Microaggressions")

add_body(doc,
    "Microaggressions are brief, everyday exchanges that communicate hostile, derogatory, "
    "or negative messages — often unintentionally. The word 'micro' does not mean the impact "
    "is small. It means the actions are subtle, but they accumulate."
)

add_title(doc, "Three Types of Microaggressions", level=2)

micro_types = [
    ("1. Microassaults", "Conscious, deliberate verbal or non-verbal attacks meant to hurt or exclude. The person committing them is usually aware of what they are doing."),
    ("2. Microinsults", "Rude or insensitive comments that demean a person's identity or heritage. Often unintentional, they convey hidden assumptions about competence, belonging, or worth."),
    ("3. Microinvalidations", "Comments or behaviours that dismiss, negate, or nullify the lived experiences of a person. These often sound positive on the surface but erase a person's reality."),
]

for title, desc in micro_types:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_TEAL
    run.font.size = Pt(11)
    add_body(doc, desc, size=10)
    add_blank_line(doc, "My example:")
    doc.add_paragraph()

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 14: MODULE 4 — VIDEO DISCUSSION & CARD SORT
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 4 — Video Discussion & Scenario Classification")

add_title(doc, "Video Discussion", level=2)
for q in ["Q1:", "Q2:", "Q3:"]:
    add_answer_block(doc, q, lines=4)

doc.add_paragraph()
add_title(doc, "Case Study Discussion — Scenario Classification", level=2)

card_rows = [
    ["1", "", ""],
    ["2", "", ""],
    ["3", "", ""],
    ["4", "", ""],
    ["5", "", ""],
    ["6", "", ""],
]
add_table(doc, ["Scenario #", "My Classification", "Correct Answer (debrief)"], card_rows)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 15: MODULE 4 — CONNECTING BACK
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 4 — Connecting Back")

add_body(doc,
    "Which bias from Module 2 drives each type of microaggression? Think about the four types "
    "of bias you explored earlier — unconscious, affinity, confirmation, and attribution. "
    "How do they connect to the microaggressions you have been identifying?"
)

for _ in range(10):
    add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 16: MODULE 5 — 3D FRAMEWORK
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 5 — Responding & Intervening")
add_title(doc, "The 3D Framework", level=2)

add_body(doc, "When you witness or experience a microaggression, you have three response options:")

ds = [
    ("Direct — Address it.", 'Speak up in the moment. Name the behaviour clearly and respectfully. Example: "That comment felt dismissive. Can we revisit what was said?"'),
    ("Distract — Redirect.", 'Shift the conversation or situation to break the pattern without direct confrontation. Example: "Hey, can we get back to the agenda?"'),
    ("Delegate — Escalate.", "Report the behaviour to someone in a position of authority, especially when there is a power imbalance or safety concern."),
]

for title, desc in ds:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_TEAL
    run.font.size = Pt(11)
    add_body(doc, desc, size=10)

doc.add_paragraph()
add_title(doc, "Choosing Your D", level=2)
add_body(doc, "The right response depends on context. All three Ds are valid.")

factor_rows = [
    ["Power dynamic", "Is there a hierarchy between you and the person?"],
    ["Severity", "How harmful was the comment or action?"],
    ["Relationship", "Do you have an established relationship with this person?"],
    ["Safety", "Is it physically and psychologically safe to speak up right now?"],
]
add_table(doc, ["Factor", "Think About"], factor_rows)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 17: MODULE 5 — QUICK CHECK & VIDEO DISCUSSION
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 5 — Quick Check & Video Discussion")

add_title(doc, "Quick Check: Which D?", level=2)
for i in range(1, 5):
    add_answer_block(doc, f"Scenario {i} — My response:", lines=2)

doc.add_paragraph()
add_title(doc, "Video Discussion", level=2)
for q in ["Q1:", "Q2:", "Q3:"]:
    add_answer_block(doc, q, lines=4)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 18: MODULE 5 — ROLE PLAY REFLECTION
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 5 — Case Study Discussion: Role-Play Reflection")

add_title(doc, "How did it feel to be:", level=2)

for role in ["The Target:", "The Person Committing the Microaggression:", "The Bystander:"]:
    add_answer_block(doc, role, lines=3)
    doc.add_paragraph()

add_answer_block(doc, "What I would do differently:", lines=4)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 19: MODULE 6 — THREE PILLARS
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 6 — From Awareness to Inclusive Action")

add_title(doc, "The Three Pillars — Connected", level=2)

add_body(doc, "Everything you have explored today forms one integrated capability:")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EI  ←→  Bias Awareness  ←→  Inclusive Action")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = DARK_TEAL

doc.add_paragraph()
connections = [
    "EI gives you the self-awareness to notice bias.",
    "Bias awareness helps you spot microaggressions.",
    "Response skills turn awareness into action.",
]
for c in connections:
    p = doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
add_title(doc, "What Inclusive Spaces Look Like", level=2)

markers = [
    "People feel safe to speak up without fear of ridicule or retaliation.",
    "Differences are valued, not merely tolerated.",
    "Feedback flows in all directions — up, down, and across.",
    "Mistakes are treated as learning opportunities, not reasons for blame.",
    "Everyone has equitable access to opportunities for growth and recognition.",
]
for i, m in enumerate(markers, 1):
    p = doc.add_paragraph(f"{i}. {m}")

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 20: MODULE 6 — TEAM COMMITMENT
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Module 6 — Team Commitment")

add_body(doc, "What is one thing your team could start doing, stop doing, or change to be more inclusive?")

doc.add_paragraph()
add_blank_line(doc, "Our team's commitment:")
for _ in range(4):
    add_blank_line(doc)

doc.add_paragraph()
add_blank_line(doc, "My role in this:")
for _ in range(3):
    add_blank_line(doc)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 21: PERSONAL ACTION PLAN
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "MY PERSONAL ACTION PLAN")

add_body(doc,
    'This is your commitment to yourself. Be specific — not "be more inclusive" but '
    '"when I hear someone being talked over in a meeting, I will use the Direct approach '
    'to redirect the conversation."',
    italic=True
)

doc.add_paragraph()

plan_items = [
    "1. One bias I will watch for in myself:",
    "2. One stereotype I will actively challenge:",
    "3. One microaggression response I will practise:",
    "4. Specific situation where I will apply this:",
    "5. How I will hold myself accountable:",
]

for item in plan_items:
    add_answer_block(doc, item, lines=2)
    doc.add_paragraph()

add_blank_line(doc, "6. My accountability partner — Name:")
add_blank_line(doc, "   Contact:")

doc.add_paragraph()
add_blank_line(doc, "7. 30-day check-in date:")

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGE 22: CLOSING
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

add_title(doc, "Closing")

add_title(doc, "My EQ Prism Comparison", level=2)

comparison_rows = [
    ["Before (Morning)", "", ""],
    ["After (Afternoon)", "", ""],
]
add_table(doc, ["", "Zone", "Notes"], comparison_rows)

doc.add_paragraph()
add_blank_line(doc, "What shifted:")
for _ in range(2):
    add_blank_line(doc)

doc.add_paragraph()
add_title(doc, "My Ongoing Toolkit", level=2)

toolkit_items = [
    "This handbook and my personal action plan",
    "My accountability partner",
    "OurEQ.LIFE — free self-awareness check-in tool (visit: OurEQ.LIFE/prism)",
]
for item in toolkit_items:
    p = doc.add_paragraph()
    run = p.add_run("☐  ")
    run.font.size = Pt(12)
    run2 = p.add_run(item)
    run2.font.size = Pt(11)

add_body(doc, "Make the EQ Prism a weekly habit. It takes 60 seconds and helps you stay connected to how you are feeling.", italic=True, size=10)

doc.add_paragraph()
add_answer_block(doc, "My One Insight from Today:", lines=3)
doc.add_paragraph()
add_answer_block(doc, "My One Commitment:", lines=3)

add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# PAGES 23-24: NOTES
# ═══════════════════════════════════════════════════════════════
for page_num in range(2):
    add_page_break(doc)
    add_title(doc, "Notes")
    for _ in range(28):
        add_blank_line(doc)
    add_confidential_footer(doc)

# ═══════════════════════════════════════════════════════════════
# BACK COVER
# ═══════════════════════════════════════════════════════════════
add_page_break(doc)

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Thank you for your participation\nin the Disha Initiative.")
run.font.size = Pt(18)
run.font.color.rgb = DARK_TEAL
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Your willingness to reflect, learn, and commit to inclusive action makes a difference — "
    "for your team, your organisation, and the people around you."
)
run.font.size = Pt(11)
run.font.color.rgb = MEDIUM_GRAY

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Reynlab Technologies Pvt Ltd")
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = DARK_GRAY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OurEQ.LIFE")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = DARK_TEAL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Your ongoing self-awareness companion")
run.font.size = Pt(10)
run.font.color.rgb = MEDIUM_GRAY
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This workbook is provided for use during facilitated sessions only. "
    "Contents are proprietary to Reynlab Technologies Pvt Ltd. "
    "Reproduction, distribution, or use outside the facilitated session "
    "is prohibited without written permission."
)
run.font.size = Pt(7)
run.font.color.rgb = LIGHT_GRAY
run.italic = True

# ─── Save ───
output_path = "/Volumes/Dev/Course_Generator/courses/ei-microaggression-disha/participant-handbook.docx"
doc.save(output_path)
print(f"Handbook saved to: {output_path}")
