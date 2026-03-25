#!/usr/bin/env python3
"""
Generate Combined 2-Day Training Program DOCX
Social & Emotional Intelligence for Inclusive Shop Floor Leadership
Merges: IMSF Course + EI Training Proposal (Dr. Sandhya Rani C)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

BRAND_BLUE = RGBColor(0x1F, 0x5C, 0x99)   # deep blue for headings
ACCENT_BLUE = RGBColor(0x2E, 0x86, 0xC1)  # mid blue for sub-labels
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)


def set_cell_bg(cell, hex_color="1F5C99"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_colored_heading(doc, text, level=1, color=BRAND_BLUE):
    p = doc.add_heading(text, level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_module_header(doc, number, title, duration, day_color=BRAND_BLUE):
    """Render a module title block with colored label."""
    p = doc.add_paragraph()
    label = p.add_run(f"Module {number}  ")
    label.bold = True
    label.font.color.rgb = day_color
    label.font.size = Pt(12)
    name = p.add_run(title)
    name.bold = True
    name.font.size = Pt(12)
    dur = p.add_run(f"   |   {duration}")
    dur.italic = True
    dur.font.size = Pt(11)
    dur.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_schedule_row(table, row_idx, time, activity, is_break=False):
    row = table.rows[row_idx]
    row.cells[0].text = time
    row.cells[1].text = activity
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                if is_break:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(1.15)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run("TRAINING PROGRAM PROPOSAL")
t.bold = True
t.font.size = Pt(22)
t.font.color.rgb = BRAND_BLUE

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub_p.add_run("Social & Emotional Intelligence\nfor Inclusive Shop Floor Leadership")
s.bold = True
s.font.size = Pt(16)

doc.add_paragraph()

tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_p.add_run("A Combined 2-Day Intensive Training Program").italic = True

doc.add_paragraph()
doc.add_paragraph()

meta_lines = [
    ("Prepared for", "Ms. Alanknanda Sen, BTC, Bangalore"),
    ("Proposal Reference", "SRC_24_030 (Combined Edition)"),
    ("Facilitator", "Dr. Sandhya Rani C"),
    ("Date", datetime.now().strftime("%B %d, %Y")),
    ("Batch Size", "25 Participants"),
]

for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.color.rgb = BRAND_BLUE
    p.add_run(value)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PROGRAM OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Program Overview", 1)

doc.add_paragraph(
    "This two-day program integrates two complementary bodies of work into a single, "
    "cohesive learning journey: the Social & Emotional Intelligence (SEI) framework "
    "developed by Dr. Sandhya Rani C, and the Inclusive Management for Shop Floor "
    "Supervisors (IMSF) curriculum. Together they address both the internal emotional "
    "competencies and the practical management behaviors that male supervisors need to "
    "lead inclusively in a manufacturing environment."
)

doc.add_paragraph()

overview_table = doc.add_table(rows=6, cols=2)
overview_table.style = "Light Grid Accent 1"

overview_data = [
    ("Duration", "2 days  |  9:00 AM – 5:00 PM each day"),
    ("Target Audience", "Male supervisors & managers, shop floor operations (2–15 yrs experience)"),
    ("Facilitator", "Dr. Sandhya Rani C — PhD, Emotional Intelligence & Behavioral Science"),
    ("Delivery Language", "English (all examples localised to Indian manufacturing context)"),
    ("Batch Size", "25 participants"),
    ("Pre-Work Required", "Individual EI Profiling — completed online, 1 week before Day 1"),
]

for idx, (field, value) in enumerate(overview_data):
    row = overview_table.rows[idx]
    row.cells[0].text = field
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PROGRAM OBJECTIVES
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Program Objectives", 1)

objectives = [
    "Provide Individual Emotional Intelligence Profiling with personalised debriefing",
    "Develop skills to manage self and others, ensuring mental and emotional wellbeing "
    "of self and team members through experiential learning",
    "Develop sensitisation to social skills while working with diverse team members "
    "through simulations of real shop floor scenarios",
    "Build respectful and inclusive management practices toward women employees in "
    "shop floor environments",
    "Apply empathy and emotional intelligence in daily workplace interactions and "
    "conflict situations",
    "Implement appropriate workplace accommodations and support systems for diverse "
    "employee needs",
    "Provide continued support to translate learnings into sustained behaviour change "
    "through post-training review sessions",
]

for obj in objectives:
    add_bullet(doc, obj)

doc.add_paragraph()
add_colored_heading(doc, "Program Goals", 2)

goals = [
    "Prepare emotionally intelligent associates equipped for current workplace challenges "
    "and future leadership opportunities",
    "Achieve measurable, observable behaviour change in daily interactions with women "
    "employees on the shop floor",
    "Reduce employee complaints and create a culture where all employees feel respected "
    "and valued",
]

for goal in goals:
    add_bullet(doc, goal)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PRE-WORK
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Pre-Work: Individual EI Profiling", 1)

doc.add_paragraph(
    "To maximise the value of Day 1, each participant completes an individual Emotional "
    "Intelligence profiling assessment before the training begins."
)

doc.add_paragraph()
add_colored_heading(doc, "How It Works", 2)

prework_steps = [
    ("1 Week Before Day 1",
     "A profiling link is emailed to each participant. The assessment takes 15–20 minutes "
     "to complete online. The EI profiling document includes personalised developmental tips."),
    ("1 Hour Before Day 1",
     "Individual EI profiles are emailed directly to each participant's work ID. "
     "Participants review their own profile before the session begins."),
    ("Day 1 — Module 1",
     "The first 45–60 minutes of Day 1 are dedicated to profile debriefing in a group "
     "setting, helping participants understand their personal strengths and development areas."),
]

for label, desc in prework_steps:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}:  ")
    r.bold = True
    r.font.color.rgb = ACCENT_BLUE
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph(
    "Note: Participants must be uninterrupted and in a private space when completing "
    "their profiling to ensure honest, reflective responses."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# DAY 1
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Day 1 — Self Mastery", 1, BRAND_BLUE)

day1_intro = doc.add_paragraph(
    "Day 1 focuses inward — helping participants understand their own emotional "
    "landscape, identify their biases and triggers, and build the foundational "
    "self-awareness and self-regulation skills that underpin all inclusive leadership."
)

doc.add_paragraph()

# Day 1 Schedule Table
add_colored_heading(doc, "Day 1 Schedule", 2)
sched1 = doc.add_table(rows=11, cols=2)
sched1.style = "Light Grid Accent 1"

# Header
sched1.rows[0].cells[0].text = "Time"
sched1.rows[0].cells[1].text = "Activity"
for cell in sched1.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, "1F5C99")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

schedule_day1 = [
    ("09:00 – 09:15", "Welcome, Housekeeping & Ground Rules", False),
    ("09:15 – 11:30", "Module 1: Program Foundation & SEI Framework", False),
    ("11:30 – 11:45", "Morning Break", True),
    ("11:45 – 13:15", "Module 2: Emotional Self-Awareness & Unconscious Bias", False),
    ("13:15 – 14:15", "Lunch Break", True),
    ("14:15 – 15:45", "Module 3: Emotional Self-Control", False),
    ("15:45 – 16:00", "Afternoon Break", True),
    ("16:00 – 17:00", "Module 4: Personal Power", False),
    ("17:00", "Day 1 Close", False),
]

for i, (time, activity, is_break) in enumerate(schedule_day1, 1):
    add_schedule_row(sched1, i, time, activity, is_break)

doc.add_paragraph()

# ── MODULE 1 ──
add_module_header(doc, 1, "Program Foundation & SEI Framework", "2 hrs 15 min")

doc.add_paragraph(
    "This opening module sets the stage for the entire two-day journey. Participants "
    "receive debriefing on their personal EI profiles, are introduced to the SEI model, "
    "and explore the neuroscience behind emotions — grounded firmly in the context of "
    "why inclusive management on the shop floor is both a business imperative and a "
    "legal obligation."
)

add_colored_heading(doc, "Content", 3)
m1_content = [
    "Welcome, introductions and establishing psychological safety ground rules",
    "Individual EI Profile Debriefing — understanding personal strengths and challenges "
    "across the SEI competency map",
    "Introduction to the Social & Emotional Intelligence (SEI) Model and its four "
    "quadrants of competencies",
    "Stories and thought experiments for each quadrant — bringing competencies to life",
    "Neuroscience of Emotions: how the brain processes emotion, and neuroplasticity — "
    "how behaviour can change with deliberate practice",
    "Why Inclusive Management Matters: the business case for respectful leadership — "
    "impact on productivity, retention, and team performance",
    "Legal framework: workplace equality obligations in India and implications for "
    "supervisors",
    "Current state reflection: acknowledging the shop floor context and challenges "
    "that brought us here",
]
for item in m1_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m1_activities = [
    "Individual profile review (silent reading + personal reflection)",
    "Group debrief: sharing one strength and one development area",
    "Thought experiment: 'What does this competency look like on our shop floor?'",
]
for item in m1_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 2 ──
add_module_header(doc, 2, "Emotional Self-Awareness & Unconscious Bias", "1 hr 30 min")

doc.add_paragraph(
    "Participants explore the first foundational SEI competency — Emotional "
    "Self-Awareness — alongside an honest examination of unconscious bias. These two "
    "topics are closely linked: understanding our own emotional patterns is the first "
    "step to recognising how bias can operate beneath conscious awareness."
)

add_colored_heading(doc, "Content", 3)
m2_content = [
    "Competency: Emotional Self-Awareness — definition, qualities of high ESA individuals",
    "Recognising common unconscious biases in workplace settings — with shop floor "
    "examples",
    "Personal bias self-assessment activity — identifying individual bias patterns",
    "Recognising personal emotional triggers in management situations",
    "5 Development Tips for building Emotional Self-Awareness: Visualisation Method, "
    "Gratitude Practice, Emotion Labelling, Body Scan, Journalling",
    "Connecting Emotional Self-Awareness with the Neuroscience of Emotions from Module 1",
]
for item in m2_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m2_activities = [
    "Personal bias self-assessment (written reflective exercise)",
    "Small group discussion: 'Which biases have I witnessed or experienced on the shop floor?'",
    "Trigger mapping: participants identify 2–3 personal emotional triggers and their "
    "typical reactions",
]
for item in m2_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 3 ──
add_module_header(doc, 3, "Emotional Self-Control", "1 hr 30 min")

doc.add_paragraph(
    "Building on self-awareness, this module develops the ability to regulate emotions "
    "— particularly under pressure. In a demanding shop floor environment, self-control "
    "is the skill that separates reactive management from considered, respectful "
    "leadership."
)

add_colored_heading(doc, "Content", 3)
m3_content = [
    "Competency: Emotional Self-Control — definition and ability to work under pressure",
    "Qualities of people with high Self-Control — what it looks like in practice",
    "The cost of low self-control in supervisory roles: impact on team trust and "
    "psychological safety",
    "Practical techniques for emotional regulation in the moment",
    "Video case studies: self-control in shop floor scenarios — debrief and discussion",
    "Development Tips for building lasting self-regulation habits",
]
for item in m3_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m3_activities = [
    "Balloon Activity — experiential exercise demonstrating the build-up and release "
    "of emotional pressure",
    "Video case study debrief: 'What did the supervisor do well / poorly? What would "
    "you do differently?'",
    "Pair practice: applying a regulation technique in a simulated difficult conversation",
]
for item in m3_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 4 ──
add_module_header(doc, 4, "Personal Power", "1 hr")

doc.add_paragraph(
    "This module closes Day 1 by anchoring participants in a sense of personal agency "
    "and capability. Personal Power — the belief in one's ability to influence outcomes "
    "— is the internal driver that sustains behavioural change beyond the training room."
)

add_colored_heading(doc, "Content", 3)
m4_content = [
    "Competency: Personal Power — definition and its role in the SEI model",
    "Qualities of people with high Personal Power — and what holds others back",
    "The link between Personal Power and accountability for behaviour change",
    "Development Tips: Achievement discussion, Level 2 techniques for building "
    "personal agency",
    "Day 1 reflection: one commitment to carry into tomorrow",
]
for item in m4_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m4_activities = [
    "Achievement discussion: participants share a moment of positive influence from "
    "their management experience",
    "Personal commitment card: one self-regulation action to practise before Day 2",
]
for item in m4_activities:
    add_bullet(doc, item)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# DAY 2
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Day 2 — Social Intelligence & Application", 1, BRAND_BLUE)

doc.add_paragraph(
    "Day 2 turns outward — applying the self-mastery built on Day 1 to real interactions "
    "with others. Participants develop empathy, communication skills, the ability to "
    "influence without authority, and conflict management capabilities, all anchored in "
    "shop floor realities. The day closes with a structured commitment to behaviour change."
)

doc.add_paragraph()

# Day 2 Schedule Table
add_colored_heading(doc, "Day 2 Schedule", 2)
sched2 = doc.add_table(rows=11, cols=2)
sched2.style = "Light Grid Accent 1"

sched2.rows[0].cells[0].text = "Time"
sched2.rows[0].cells[1].text = "Activity"
for cell in sched2.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, "1F5C99")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

schedule_day2 = [
    ("09:00 – 09:15", "Day 1 Recap & Day 2 Overview", False),
    ("09:15 – 11:15", "Module 5: Empathy in Practice", False),
    ("11:15 – 11:30", "Morning Break", True),
    ("11:30 – 12:30", "Module 6: Professional Communication & Daily Interactions", False),
    ("12:30 – 13:30", "Lunch Break", True),
    ("13:30 – 15:00", "Module 7: Influencing Without Authority", False),
    ("15:00 – 15:15", "Afternoon Break", True),
    ("15:15 – 16:15", "Module 8: Conflict Management & Health Sensitivity", False),
    ("16:15 – 16:45", "Module 9: Application & Integration", False),
    ("16:45 – 17:00", "Module 10: Commitment & Action Planning — Closing", False),
]

for i, (time, activity, is_break) in enumerate(schedule_day2, 1):
    add_schedule_row(sched2, i, time, activity, is_break)

doc.add_paragraph()

# ── MODULE 5 ──
add_module_header(doc, 5, "Empathy in Practice", "2 hrs")

doc.add_paragraph(
    "Empathy is the bridge between self-mastery and effective social leadership. This "
    "module combines the SEI Empathy competency with a direct focus on understanding "
    "women's experiences in the shop floor environment — making the learning immediately "
    "relevant and impactful."
)

add_colored_heading(doc, "Content", 3)
m5_content = [
    "Competency: Empathy — definition, neurological basis, and why it matters in leadership",
    "Qualities of people with high empathy — and what low empathy costs a team",
    "Women's perspectives on shop floor challenges — video testimonials (anonymous, "
    "real workplace experiences)",
    "Empathy vs sympathy: understanding the critical distinction",
    "Empathy mapping: walking in another person's shoes through a structured framework",
    "Active empathetic listening techniques — practice in pairs",
    "Empathy in non-face-to-face settings: virtual check-ins, WhatsApp messages, "
    "phone calls",
    "Development Tips for building and sustaining empathy as a daily practice",
]
for item in m5_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m5_activities = [
    "Video testimonial viewing and group debrief: 'What did you hear? What surprised you? "
    "What will you do differently?'",
    "Empathy mapping exercise: participants map the experience of a woman colleague in "
    "a specific shop floor scenario",
    "Paired listening practice: active empathetic listening with structured feedback",
]
for item in m5_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 6 ──
add_module_header(doc, 6, "Professional Communication & Daily Interactions", "1 hr")

doc.add_paragraph(
    "Translating empathy into concrete, everyday behaviour — this module gives "
    "participants the specific communication tools and scripts they need to interact "
    "with women employees professionally and respectfully from the very next working day."
)

add_colored_heading(doc, "Content", 3)
m6_content = [
    "Professional greetings and daily courtesies in a shop floor context",
    "Respectful vs disrespectful behaviours — live demonstration and analysis",
    "Language that includes and language that excludes: practical examples",
    "Delivering constructive feedback to women employees professionally and without bias",
    "Handling sensitive requests with dignity (seating, scheduling, workload)",
    "Do's and Don'ts: a shop floor communication reference guide",
]
for item in m6_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m6_activities = [
    "Demonstration: facilitator models respectful and disrespectful interaction styles — "
    "group identifies the differences",
    "Role-play 1: delivering constructive feedback to a woman employee",
    "Role-play 2: responding appropriately to a sensitive workplace request",
]
for item in m6_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 7 ──
add_module_header(doc, 7, "Influencing Without Authority", "1 hr 30 min")

doc.add_paragraph(
    "Effective inclusive leadership often requires influencing peers, team members, "
    "and even upward — without relying on positional power. This module equips "
    "participants with practical influencing skills grounded in the SEI framework."
)

add_colored_heading(doc, "Content", 3)
m7_content = [
    "Competency: Influencing Others Without Power — definition and relevance to "
    "shop floor leadership",
    "Qualities of people with high ability to influence others",
    "The difference between influence and manipulation — ethical boundaries",
    "Influencing techniques: building credibility, finding common ground, "
    "storytelling, reciprocity",
    "Applying influence in the context of creating an inclusive team culture",
    "Video case study: influence in action — debrief and lessons for the shop floor",
    "Development Tips",
]
for item in m7_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m7_activities = [
    "Lost at Sea — group simulation exercise revealing natural influencing styles "
    "and team dynamics",
    "Group debrief: 'Who influenced the group's decisions and how?'",
    "Video case study debrief: identifying effective and ineffective influence tactics",
]
for item in m7_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 8 ──
add_module_header(doc, 8, "Conflict Management & Health Sensitivity", "1 hr")

doc.add_paragraph(
    "This module addresses two interconnected realities of inclusive shop floor "
    "management: handling conflict with professionalism, and responding sensitively "
    "to health-related needs — particularly those specific to women employees. Both "
    "require the self-control and empathy built in earlier modules."
)

add_colored_heading(doc, "Content", 3)
m8_content = [
    "Competency: Conflict Management — definition and the five conflict management styles",
    "Self-assessment of personal conflict management style",
    "The Conflict GRID: mapping styles to outcomes with shop floor workplace examples",
    "When and how to use each conflict style appropriately",
    "Understanding menstrual health needs in the workplace context — facts, myths, "
    "and the supervisor's role",
    "Responding professionally to health-related accommodation requests (seating, "
    "breaks, schedule adjustments) with dignity and discretion",
    "Proposing team norms that create a naturally supportive environment for all "
    "employees",
    "Development Tips for conflict management",
]
for item in m8_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m8_activities = [
    "Conflict style self-assessment — participants identify their dominant style",
    "Scenario analysis: applying the GRID to three real shop floor conflict situations",
    "Group discussion: co-creating a list of team norms for a supportive environment",
    "Quick quiz on conflict management (5 questions)",
]
for item in m8_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 9 ──
add_module_header(doc, 9, "Application & Integration", "30 min")

doc.add_paragraph(
    "Participants bring together everything from both days through structured case "
    "study analysis. This module reinforces learning transfer and helps consolidate "
    "the connection between SEI competencies and inclusive management behaviours in "
    "a shop floor setting."
)

add_colored_heading(doc, "Content", 3)
m9_content = [
    "Case study analysis using inclusive management frameworks (2–3 shop floor "
    "scenarios covering bias, conflict, health accommodation, and daily interaction)",
    "Small group analysis: applying SEI competencies to diagnose and resolve each scenario",
    "Group presentations: teams share recommended responses and rationale",
    "Facilitator synthesis: key takeaways and connecting threads across both days",
]
for item in m9_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m9_activities = [
    "Small group case study analysis (3 groups, 1 scenario each)",
    "Group presentations and peer feedback",
]
for item in m9_activities:
    add_bullet(doc, item)

doc.add_paragraph()

# ── MODULE 10 ──
add_module_header(doc, 10, "Commitment & Action Planning  —  Closing", "15 min")

doc.add_paragraph(
    "The program closes with each participant making a personal, public commitment "
    "to specific behaviour changes. This accountability mechanism significantly "
    "increases the likelihood of sustained change beyond the training room."
)

add_colored_heading(doc, "Content", 3)
m10_content = [
    "Personal Action Plan: each participant identifies 3 specific behavioural "
    "commitments they will implement in the next 30 days",
    "Commitments are written using the format: 'By [date], I will [specific behaviour] "
    "when [specific situation]'",
    "Public Commitment Circle: each participant shares one commitment with the group",
    "Overview of post-training review schedule (Day 7 and Day 21)",
    "Closing remarks and next steps",
]
for item in m10_content:
    add_bullet(doc, item)

add_colored_heading(doc, "Activities", 3)
m10_activities = [
    "Personal Action Plan completion (individual, written)",
    "Commitment Circle: each participant reads one commitment aloud",
]
for item in m10_activities:
    add_bullet(doc, item)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# POST-TRAINING REVIEW SESSIONS
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Post-Training Review Sessions", 1)

doc.add_paragraph(
    "Behaviour change does not happen in a training room — it happens in the days "
    "and weeks that follow. Two structured review sessions are scheduled for each "
    "cohort to sustain momentum, address challenges, and reinforce commitments."
)

doc.add_paragraph()

review_table = doc.add_table(rows=3, cols=3)
review_table.style = "Light Grid Accent 1"

review_table.rows[0].cells[0].text = "Session"
review_table.rows[0].cells[1].text = "Timing"
review_table.rows[0].cells[2].text = "Purpose & Format"
for cell in review_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

review_data = [
    ("Review Session 1", "Day 7 post-training\n(1 hour per cohort)",
     "Check-in on personal action plans • Share early wins and obstacles • "
     "Peer support and accountability • Facilitator guidance on sticking points"),
    ("Review Session 2", "Day 21 post-training\n(1 hour per cohort)",
     "Progress review against action plan commitments • Debrief of real situations "
     "handled using program tools • Celebrate behaviour changes • "
     "Set commitments for the next 30 days"),
]

for idx, (session, timing, purpose) in enumerate(review_data, 1):
    review_table.rows[idx].cells[0].text = session
    review_table.rows[idx].cells[1].text = timing
    review_table.rows[idx].cells[2].text = purpose

doc.add_paragraph()
doc.add_paragraph(
    "Progress from review sessions is reported by participants and tracked against "
    "the original action plans. Patterns and themes are shared (anonymously) with "
    "HR leadership to inform ongoing support needs."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Methodology", 1)

add_colored_heading(doc, "Instructional Approach", 2)

doc.add_paragraph(
    "The program uses experiential learning methods throughout — participants learn "
    "by doing, reflecting, and practising, not by listening to lectures. All content "
    "is grounded in the neuroscience of behaviour change and contextualised for the "
    "Indian manufacturing shop floor."
)

methods = [
    ("Video Testimonials & Case Studies",
     "Real-life examples from workplace settings, using anonymous cases and "
     "storytelling methods to build empathy and contextual understanding"),
    ("Simulations & Group Activities",
     "Structured simulations (Lost at Sea, Balloon Activity, PIPE) that surface "
     "natural behavioural patterns and create memorable learning moments"),
    ("Role-Play Exercises",
     "Structured practice of key communication skills with peer feedback, allowing "
     "participants to rehearse new behaviours in a safe environment"),
    ("Reflective Exercises",
     "Individual and paired reflections (trigger mapping, empathy mapping, bias "
     "self-assessment) that deepen self-awareness"),
    ("Small Group Discussions",
     "Peer learning through structured group discussions, leveraging the collective "
     "experience in the room"),
    ("Personal Action Planning",
     "Structured commitment tool ensuring each participant leaves with specific, "
     "measurable behaviour change goals"),
]

for title, desc in methods:
    p = doc.add_paragraph()
    p.add_run(f"{title}:  ").bold = True
    p.add_run(desc)

doc.add_paragraph()
add_colored_heading(doc, "Participant Requirements", 2)

participant_reqs = [
    "Complete EI Profiling online 1 week before Day 1 (15–20 minutes)",
    "Full attendance across both days (9:00 AM – 5:00 PM)",
    "Openness to self-reflection and active participation in all activities",
    "Participation in Day 7 and Day 21 review sessions",
    "Completion of personal action plan at end of Day 2",
]
for req in participant_reqs:
    add_bullet(doc, req)

doc.add_paragraph()
add_colored_heading(doc, "Materials Provided", 2)

materials = [
    "Individual EI Profile report (personalised, pre-session)",
    "Participant workbook (printed) — includes session notes, activity worksheets, "
    "homework sheets and progress tracking pages",
    "Personal Action Plan template",
    "Quick Reference Card: SEI competencies and daily practice tips",
    "Video content (empathy, self-control, influence, unconscious bias)",
    "Case study materials",
]
for mat in materials:
    add_bullet(doc, mat)

doc.add_paragraph()
add_colored_heading(doc, "Logistics", 2)

logistics = [
    ("Batch Size", "25 participants"),
    ("Room Setup", "U-shape seating configuration"),
    ("Equipment", "Projector / screen, sound system, flipcharts and markers"),
    ("Location", "Training facility to be provided by the organisation"),
    ("Travel & Accommodation",
     "To be provided by the organisation (not required for Coimbatore location)"),
]

for label, value in logistics:
    p = doc.add_paragraph()
    p.add_run(f"{label}:  ").bold = True
    p.add_run(value)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# ASSESSMENT
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Assessment", 1)

add_colored_heading(doc, "EI Profiling (Pre-Training Baseline)", 2)

doc.add_paragraph(
    "Each participant completes an Individual Emotional Intelligence Profiling "
    "assessment before Day 1. This serves as both a personalised learning tool "
    "and a baseline measurement of SEI competency levels across the cohort. "
    "The profile includes developmental tips tailored to each individual's results."
)

doc.add_paragraph()
add_colored_heading(doc, "In-Training Assessment", 2)

in_training = [
    ("Active Participation",
     "Engagement in all activities, discussions, simulations, and role-plays "
     "across both days"),
    ("Conflict Management Quiz",
     "5-question quiz at end of Module 8 to consolidate understanding of "
     "conflict management styles and approaches"),
    ("Case Study Analysis",
     "Group analysis and presentation during Module 9 — demonstrating ability "
     "to apply SEI competencies to real shop floor scenarios"),
    ("Personal Action Plan",
     "Completion of a specific, measurable 30-day behaviour change plan at "
     "the close of Day 2"),
]

for label, desc in in_training:
    p = doc.add_paragraph()
    p.add_run(f"{label}:  ").bold = True
    p.add_run(desc)

doc.add_paragraph()
add_colored_heading(doc, "Post-Training Impact Tracking", 2)

post_training = [
    "Action plan progress reported at Day 7 and Day 21 review sessions",
    "Reduction in employee complaints tracked at 30, 60, and 90 days",
    "Anonymous feedback from women employees on observed supervisor behaviour changes",
    "Management observation of behavioural improvements in daily operations",
]
for item in post_training:
    add_bullet(doc, item)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# FACILITATOR PROFILE
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Facilitator Profile", 1)

doc.add_paragraph(
    "Dr. Sandhya Rani C is a behavioural scientist and emotional intelligence "
    "specialist with a PhD in Emotional Intelligence. She brings scientific rigour, "
    "deep practitioner experience, and an authentic gender perspective to every "
    "session she facilitates."
)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("LinkedIn:  ").bold = True
p.add_run("https://www.linkedin.com/in/dr-sandhya-rani-c/")

doc.add_paragraph()
add_colored_heading(doc, "Why an External Female Facilitator?", 2)

rationale = [
    ("Scientific Credibility",
     "PhD-level expertise establishes authority on empathy, bias, and emotional "
     "intelligence, ensuring participants take the content seriously"),
    ("Authentic Gender Perspective",
     "A female facilitator brings lived perspective on women's workplace experiences, "
     "giving the training authenticity that cannot be replicated otherwise"),
    ("Psychological Safety",
     "As an external consultant, Dr. Sandhya creates a neutral, non-judgemental "
     "space that allows participants to be genuinely reflective"),
    ("Modelling the Behaviour",
     "Participants experience a respectful, professional interaction with an "
     "accomplished woman leader — the training in action"),
    ("Behavioural Change Expertise",
     "Deep background in habit formation, neuroplasticity, and organisational "
     "behaviour change ensures the program is designed for lasting impact"),
]

for title, desc in rationale:
    p = doc.add_paragraph()
    p.add_run(f"{title}:  ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# BUDGET
# ─────────────────────────────────────────────────────────────────────────────

add_colored_heading(doc, "Budget", 1)

budget_table = doc.add_table(rows=4, cols=3)
budget_table.style = "Light Grid Accent 1"

budget_table.rows[0].cells[0].text = "Session"
budget_table.rows[0].cells[1].text = "Cost (INR)"
budget_table.rows[0].cells[2].text = "Total (INR)"
for cell in budget_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, "1F5C99")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

budget_data = [
    ("Social & Emotional Intelligence Training\n(2-Day Combined Program)", "2,02,000", "2,02,000"),
    ("Review Session — Day 7 (per cohort)", "Included", "—"),
    ("Review Session — Day 21 (per cohort)", "Included", "—"),
]

for idx, (session, cost, total) in enumerate(budget_data, 1):
    budget_table.rows[idx].cells[0].text = session
    budget_table.rows[idx].cells[1].text = cost
    budget_table.rows[idx].cells[2].text = total

doc.add_paragraph()
doc.add_paragraph("* GST extra")
doc.add_paragraph(
    "Note: Review sessions on Day 7 and Day 21 are included within the program fee. "
    "Travel and accommodation for the facilitator to be provided by the organisation "
    "(not applicable for Coimbatore location)."
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

output_path = (
    "/Volumes/Dev/Course_Generator/courses/inclusive-management-shop-floor/"
    "SEI_IMSF_Combined_Training_Proposal.docx"
)
doc.save(output_path)
print(f"Document saved: {output_path}")
