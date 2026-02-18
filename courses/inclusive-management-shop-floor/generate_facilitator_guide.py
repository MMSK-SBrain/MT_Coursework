"""
generate_facilitator_guide.py
Generates the Pre/Post Self-Assessment Facilitator Guide .docx file.

Usage:
    python3 generate_facilitator_guide.py

Output:
    assessments/pre-post-self-assessment-facilitator-guide.docx
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Path resolution
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSESSMENTS_DIR = os.path.join(SCRIPT_DIR, "assessments")
os.makedirs(ASSESSMENTS_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(ASSESSMENTS_DIR, "pre-post-self-assessment-facilitator-guide.docx")

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def apply_shading(cell, fill_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_para_shading(paragraph, fill_hex):
    """Apply background shading to a paragraph (for cover header)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_script_box(doc, text):
    """Add spoken text in a shaded single-cell table box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    apply_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return tbl


def add_stage_direction(doc, text):
    """Add a stage direction paragraph in bold italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_body_para(doc, text, italic=False, bold=False, size=11):
    """Add a regular body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_heading(doc, text, level=1):
    """Add a Heading paragraph."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
    return p


def set_cell_text_bold(cell, text, size=11):
    """Set cell text as bold."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def set_cell_text(cell, text, bold=False, size=11):
    """Set cell text with optional bold."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def create_document():
    doc = Document()

    # Page setup: A4, 2cm margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Default font via Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    return doc, section


# ---------------------------------------------------------------------------
# Page header (runs on every page)
# ---------------------------------------------------------------------------

def add_page_header(section):
    """Add 'FACILITATOR GUIDE -- NOT FOR DISTRIBUTION' to every page header."""
    header = section.header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("FACILITATOR GUIDE -- NOT FOR DISTRIBUTION")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# COVER SECTION
# ---------------------------------------------------------------------------

def build_cover(doc):
    # "FACILITATOR GUIDE -- NOT FOR DISTRIBUTION" — shaded, bold, 14pt, centered
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("FACILITATOR GUIDE -- NOT FOR DISTRIBUTION")
    run1.bold = True
    run1.font.size = Pt(14)
    run1.font.name = "Calibri"
    set_para_shading(p1, "D9D9D9")  # light grey

    doc.add_paragraph()  # blank line

    # Title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Pre/Post Self-Assessment Administration Guide")
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.name = "Calibri"

    # Subtitle
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Inclusive Management for Shop Floor Supervisors")
    run3.italic = True
    run3.font.size = Pt(12)
    run3.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Guidance note
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(
        "This guide contains verbatim scripts, scoring information, and administration protocols."
    )
    run4.font.size = Pt(10)
    run4.font.name = "Calibri"

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Do not distribute to participants.")
    run5.bold = True
    run5.font.size = Pt(10)
    run5.font.name = "Calibri"

    # Page break
    doc.add_page_break()

    print("PASS: Cover/Header")


# ---------------------------------------------------------------------------
# SECTION 1: Item-Domain Map
# ---------------------------------------------------------------------------

# All 21 items with: (number, domain, first ~8 words + "...", direction)
ITEMS = [
    (1,  "Foundation",           "I treat every person on my team with...",                   "Forward"),
    (2,  "Foundation",           "When I make a decision that affects my...",                  "Forward"),
    (3,  "Foundation",           "I hold myself to the same behavioral...",                    "Forward"),
    (4,  "Daily Interactions",   "When a woman on my team speaks during...",                   "Forward"),
    (5,  "Daily Interactions",   "When I need someone to take on an extra...",                 "(R)"),
    (6,  "Daily Interactions",   "I greet all team members at the start...",                   "Forward"),
    (7,  "Self-Awareness",       "Before reacting to a situation that frustrates...",          "Forward"),
    (8,  "Self-Awareness",       "I notice when my tone or manner changes...",                 "Forward"),
    (9,  "Self-Awareness",       "I regularly consider whether my assumptions about...",       "Forward"),
    (10, "Empathy",              "When a team member tells me about a problem...",             "Forward"),
    (11, "Empathy",              "I find it easier to discuss personal or work...",            "(R)"),
    (12, "Empathy",              "When a woman on my team seems upset or...",                  "Forward"),
    (13, "Feedback & Conflict",  "When I give corrective feedback, I focus on...",             "Forward"),
    (14, "Feedback & Conflict",  "I give women on my team the same quality...",               "Forward"),
    (15, "Feedback & Conflict",  "When there is a conflict on my team, I hear...",            "Forward"),
    (16, "Health Sensitivity",   "If a team member appears unwell or in physical...",          "Forward"),
    (17, "Health Sensitivity",   "When a team member requests time off or a schedule...",      "Forward"),
    (18, "Health Sensitivity",   "I hesitate to make task adjustments for a team...",          "(R)"),
    (19, "Team Inclusion",       "I assign high-visibility or challenging tasks to women...",  "Forward"),
    (20, "Team Inclusion",       "When I plan team activities or recognize team achievements...", "Forward"),
    (21, "Team Inclusion",       "I actively create conditions where every team member...",    "Forward"),
]


def build_section1(doc):
    add_heading(doc, "Section 1: Item-Domain Map", level=1)

    add_body_para(
        doc,
        "The 21 items are distributed across 7 domains (3 items each). Use this map when "
        "reviewing individual or aggregate responses, or when identifying which domain a "
        "participant's pattern of ratings addresses. Domain names do NOT appear on the "
        "participant form — visual separators only."
    )

    # Table: Item # | Domain | Item Text | Direction
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    # Set column widths
    widths = [Cm(1.5), Cm(4.5), Cm(8.5), Cm(2.5)]
    for i, w in enumerate(widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    # Header row
    headers = ["Item #", "Domain", "Item Text (abbreviated)", "Direction"]
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        set_cell_text_bold(hrow.cells[i], h, size=10)
        apply_shading(hrow.cells[i], "D9D9D9")

    # Data rows
    for item_num, domain, abbrev, direction in ITEMS:
        row = tbl.add_row()
        set_cell_text(row.cells[0], str(item_num), size=10)
        set_cell_text(row.cells[1], domain, size=10)
        set_cell_text(row.cells[2], abbrev, size=10)
        is_reverse = direction == "(R)"
        set_cell_text(row.cells[3], direction, bold=is_reverse, size=10)
        if is_reverse:
            apply_shading(row.cells[3], "FFF2CC")  # light yellow highlight for reverse items

    print("PASS: Section 1 Domain Map")


# ---------------------------------------------------------------------------
# SECTION 2: Reverse-Scored Items
# ---------------------------------------------------------------------------

REVERSE_ITEMS = {
    5: ("Daily Interactions",
        "When I need someone to take on an extra task quickly, I turn to male team members "
        "before considering women on my team."),
    11: ("Empathy",
         "I find it easier to discuss personal or work difficulties with male team members "
         "than with women on my team."),
    18: ("Health Sensitivity",
         "I hesitate to make task adjustments for a team member's health issues because I "
         "worry it will seem like preferential treatment."),
}


def build_section2(doc):
    add_heading(doc, "Section 2: Reverse-Scored Items", level=1)

    add_body_para(
        doc,
        "Three items are reverse-scored. On these items, a higher participant rating indicates "
        "less inclusive behavior. When calculating domain or total scores, apply the conversion "
        "table below before summing. The participant form carries NO indication of which items "
        "are reverse-scored."
    )

    doc.add_paragraph()

    for item_num, (domain, full_text) in REVERSE_ITEMS.items():
        p = doc.add_paragraph()
        run = p.add_run(f"Item {item_num} (R) — {domain}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"

        add_body_para(doc, full_text, italic=True)
        doc.add_paragraph()

    # Conversion table
    p_heading = doc.add_paragraph()
    run_h = p_heading.add_run("Reverse-Score Conversion Table")
    run_h.bold = True
    run_h.font.size = Pt(11)
    run_h.font.name = "Calibri"

    rs_tbl = doc.add_table(rows=6, cols=2)
    rs_tbl.style = "Table Grid"

    header_data = ["Participant marked", "Score as"]
    for i, h in enumerate(header_data):
        set_cell_text_bold(rs_tbl.rows[0].cells[i], h, size=11)
        apply_shading(rs_tbl.rows[0].cells[i], "D9D9D9")

    conversion = [("1", "5"), ("2", "4"), ("3", "3"), ("4", "2"), ("5", "1")]
    for i, (marked, score) in enumerate(conversion, 1):
        set_cell_text(rs_tbl.rows[i].cells[0], marked, size=11)
        set_cell_text(rs_tbl.rows[i].cells[1], score, size=11)

    doc.add_paragraph()

    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        "Important: The participant form carries NO indication of which items are reverse-scored. "
        "Do not mention reverse scoring to participants."
    )
    run_note.bold = True
    run_note.font.size = Pt(11)
    run_note.font.name = "Calibri"

    print("PASS: Section 2 Reverse-Scoring")


# ---------------------------------------------------------------------------
# SECTION 3: Pre-Assessment Verbatim Script
# ---------------------------------------------------------------------------

PRE_SCRIPT_INTRO = (
    "Delivery context: 9:00 AM, before any training content. Participants have just sat down. "
    "Read every word of the script below exactly as written. Do not paraphrase, explain further, "
    "or improvise."
)

PRE_SCRIPT_MAIN = (
    "Before we begin today's content, I'd like to take five minutes for a quick personal "
    "reflection exercise.\n\n"
    "Please look at the form in front of you. It's called 'My Management Practices.' Take a "
    "moment to read the header.\n\n"
    "[Pause 15 seconds.]\n\n"
    "A few important things about this form. First: I will not collect it. At the end of the "
    "day, this form goes home with you. Nobody else will see your responses — not HR, not your "
    "manager, not me. It is for your eyes only.\n\n"
    "Second: there are no right or wrong answers on this form. The only valuable answer is an "
    "honest one.\n\n"
    "Third: please rate each statement based on what you actually do in a typical work situation "
    "— not what you think you should do, and not what you intend to do. The behavior you have "
    "actually shown in the past month.\n\n"
    "You'll see two columns on the right — Pre (AM) and Post (PM). Please only fill in the Pre "
    "column right now. The Post column you'll complete at the end of our day together.\n\n"
    "For each statement, circle the number that best describes how often you actually do that "
    "behavior: 1 means Almost Never, 5 means Almost Always.\n\n"
    "Please read each statement carefully before circling.\n\n"
    "You have five minutes. Please work silently and independently.\n\n"
    "[Start timer. Remain at the front or seated. Do not circulate. Do not look at participant forms.]"
)

PRE_CONTINGENCY_WORD = (
    "[If a participant asks what a word means:]\n"
    "\"[Restate the item in simpler language / offer the Hindi equivalent if prepared.] "
    "Rate what feels most accurate to you.\""
)

PRE_CONTINGENCY_GRADED = (
    "[If a participant asks if the form is graded:]\n"
    "\"It is not graded. No one scores this except you, and you don't have to share your "
    "ratings with anyone.\""
)

PRE_CLOSING = (
    "[After five minutes:]\n"
    "\"Please set down your pen. You'll come back to the Post column at the end of our session.\""
)


def build_section3(doc):
    add_heading(doc, "Section 3: Pre-Assessment Verbatim Script", level=1)

    add_body_para(doc, PRE_SCRIPT_INTRO, italic=True)
    doc.add_paragraph()

    add_stage_direction(doc, "[Distribute forms. Wait until everyone has one. Then say:]")
    add_script_box(doc, PRE_SCRIPT_MAIN)
    doc.add_paragraph()

    add_stage_direction(doc, "[Contingency responses — read verbatim if needed:]")
    doc.add_paragraph()

    add_script_box(doc, PRE_CONTINGENCY_WORD)
    doc.add_paragraph()

    add_script_box(doc, PRE_CONTINGENCY_GRADED)
    doc.add_paragraph()

    add_script_box(doc, PRE_CLOSING)

    print("PASS: Section 3 Pre Script")


# ---------------------------------------------------------------------------
# SECTION 4: Post-Assessment Verbatim Script
# ---------------------------------------------------------------------------

POST_SCRIPT_INTRO = (
    "Delivery context: 4:45 PM, after all training content. Participants are tired but have "
    "completed the day. Read every word exactly as written."
)

POST_SCRIPT_MAIN = (
    "We've reached the final activity of the day. Please find your 'My Management Practices' "
    "form from this morning.\n\n"
    "You'll see your ratings in the Pre (AM) column — the answers you gave at 9 o'clock. Now "
    "I'd like you to complete the Post (PM) column.\n\n"
    "Please read each statement again. For each one, rate yourself based on how you see things "
    "right now — at this moment, at the end of today.\n\n"
    "It is completely fine to give the same rating as this morning. It is fine to give a higher "
    "rating. It is also fine to give a lower rating — sometimes a day of learning makes us "
    "realize we were overrating ourselves earlier, and that honest recognition is exactly what "
    "this tool is for.\n\n"
    "Rate what is true for you right now, not what you hope will be true tomorrow.\n\n"
    "[Pause.]\n\n"
    "After you've completed the Post column, please turn to the Reflection section at the bottom "
    "and answer those three questions for yourself. This is private — you don't need to share "
    "your answers.\n\n"
    "You have five minutes for the Post ratings and the reflection questions.\n\n"
    "[Start timer. Remain seated at the front or step out of the room. Do not circulate during "
    "post-assessment. This reduces social desirability pressure during the most honest part of "
    "the day.]"
)

POST_CLOSING = (
    "[After five minutes:]\n"
    "\"Thank you. In a moment, I'll share a couple of questions about what you noticed — but "
    "you won't need to share your specific ratings. Please keep your form.\""
)


def build_section4(doc):
    add_heading(doc, "Section 4: Post-Assessment Verbatim Script", level=1)

    add_body_para(doc, POST_SCRIPT_INTRO, italic=True)
    doc.add_paragraph()

    add_stage_direction(
        doc,
        "[Participants should have their form from the morning. Redistribute if any have been "
        "set aside. Then say:]"
    )
    add_script_box(doc, POST_SCRIPT_MAIN)
    doc.add_paragraph()

    add_script_box(doc, POST_CLOSING)

    print("PASS: Section 4 Post Script")


# ---------------------------------------------------------------------------
# SECTION 5: Show-of-Hands Protocol
# ---------------------------------------------------------------------------

SOH_INTRO_TEXT = (
    "This happens immediately after the post-assessment, before the Action Planning module. "
    "It takes 3-5 minutes. The facilitator reads questions aloud; participants raise hands. "
    "No individual is identified. The facilitator uses responses to gauge overall shift and "
    "frame the action planning discussion."
)

SOH_INTRODUCTION = (
    "I'd like to take a quick moment to understand what today meant for the group — not "
    "individual ratings, just a general sense of the room. Raise your hand if the answer "
    "applies to you. There are no right or wrong hands to raise."
)

SOH_QUESTIONS = [
    (
        "Question 1 — Overall shift:",
        "How many of you noticed at least one item where your Post rating was higher than "
        "your Pre rating — even by just one point?",
        "[Wait for hands. Note rough proportion. Say:] \"Good. That's awareness in action.\""
    ),
    (
        "Question 2 — Surprising self-discovery:",
        "How many of you had at least one item that surprised you — either because you rated "
        "yourself lower than you expected, or because you realized you weren't sure of your answer?",
        "[Wait for hands.] \"That uncertainty is actually a sign of honest reflection.\""
    ),
    (
        "Question 3 — Health and accommodation items (Group 6):",
        "How many of you found the items about responding to team members' health situations "
        "— items 16 to 18 — the most thought-provoking group?",
        "[Wait for hands. This flags engagement with the most sensitive domain without naming "
        "menstrual health.]"
    ),
    (
        "Question 4 — Bridge to action planning:",
        "How many of you already have a sense of one specific behavior you want to work on "
        "— something concrete you can do differently next week?",
        "[Wait for hands. Then say:] \"For those who raised their hands — that's what we're "
        "about to do in Module 7. For those who are still thinking — that's exactly what "
        "Module 7 is designed to help with. Let's move to the action planning.\""
    ),
]

SOH_NOTES = [
    "Never ask: 'How many of you gave yourself a low score?' — this triggers shame.",
    "Never identify anyone: 'I see Rajesh has his hand up...'",
    "If very few hands go up on Question 1 (possible ceiling effect at pre-assessment): say "
    "'That can mean the day confirmed what you already knew — which is also useful data.'",
    "The show of hands is for facilitator room-reading only. Do not write down hand counts "
    "in front of participants.",
]


def build_section5(doc):
    add_heading(doc, "Section 5: Show-of-Hands Protocol", level=1)

    add_body_para(doc, SOH_INTRO_TEXT, italic=True)
    doc.add_paragraph()

    add_stage_direction(doc, "[Introduce:]")
    add_script_box(doc, SOH_INTRODUCTION)
    doc.add_paragraph()

    for q_label, q_text, q_response in SOH_QUESTIONS:
        p_label = doc.add_paragraph()
        run_label = p_label.add_run(q_label)
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.name = "Calibri"

        p_q = doc.add_paragraph()
        run_q = p_q.add_run(q_text)
        run_q.bold = True
        run_q.font.size = Pt(11)
        run_q.font.name = "Calibri"

        add_body_para(doc, q_response)
        doc.add_paragraph()

    # Facilitator notes
    p_notes_head = doc.add_paragraph()
    run_nh = p_notes_head.add_run("Facilitator Notes for Show of Hands:")
    run_nh.bold = True
    run_nh.font.size = Pt(11)
    run_nh.font.name = "Calibri"

    for note in SOH_NOTES:
        p_note = doc.add_paragraph(style="List Bullet")
        run_note = p_note.add_run(note)
        run_note.font.size = Pt(11)
        run_note.font.name = "Calibri"

    print("PASS: Section 5 Show-of-Hands")


# ---------------------------------------------------------------------------
# SECTION 6: Anonymous Tear-Off Strip
# ---------------------------------------------------------------------------

TEAROFF_DESIGN = (
    "The tear-off strip is an optional feature for facilitators who want aggregate quantitative "
    "data. It appears at the very bottom of the participant form, separated by a dashed line "
    "and scissors icon. The strip reads:\n\n"
    "\"[scissors] - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -\n"
    "OPTIONAL: Anonymous Data Strip — Tear off and drop in the box at the front. "
    "Do not write your name.\n\n"
    "My Pre (AM) total score (add all 21 Pre ratings):   _______  / 105\n"
    "My Post (PM) total score (add all 21 Post ratings): _______  / 105\n"
    "My biggest single-item shift (how many points?):    _______  points\""
)

TEAROFF_CONDITIONS = [
    "She wants to track program-level pre/post data across multiple sessions.",
    "The client organization has requested aggregate metrics.",
    "She judges the group's trust level is high enough that the strip does not undermine "
    "the privacy narrative.",
]

TEAROFF_SCRIPT = (
    "This optional strip at the bottom of your form is for those who want to contribute "
    "anonymously to our program data. There is no name on it. You do not have to use it. "
    "If you'd like to contribute, tear it off and drop it in the box near the door on your "
    "way out. Most people take the whole form home."
)


def build_section6(doc):
    add_heading(doc, "Section 6: Anonymous Tear-Off Strip", level=1)

    add_body_para(doc, TEAROFF_DESIGN)
    doc.add_paragraph()

    add_heading(doc, "When to Use", level=2)

    add_body_para(
        doc,
        "The facilitator should use the tear-off strip only if all of the following apply:"
    )
    for condition in TEAROFF_CONDITIONS:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(condition)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()
    p_doubt = doc.add_paragraph()
    run_doubt = p_doubt.add_run("If in doubt, omit it.")
    run_doubt.bold = True
    run_doubt.font.size = Pt(11)
    run_doubt.font.name = "Calibri"
    add_body_para(
        doc,
        "The show-of-hands protocol provides adequate room-level data for facilitation purposes."
    )

    doc.add_paragraph()
    add_heading(doc, "How to Position", level=2)

    add_stage_direction(doc, "[Read verbatim to participants:]")
    add_script_box(doc, TEAROFF_SCRIPT)

    print("PASS: Section 6 Tear-off Strip")


# ---------------------------------------------------------------------------
# SECTION 7: Cultural Context Notes
# ---------------------------------------------------------------------------

HEALTH_SENSITIVITY_NOTE = (
    "Items 16, 17, and 18 use health-general language deliberately. This is not avoidance "
    "— it is sequencing. The pre-assessment happens before Module 5, which addresses menstrual "
    "health explicitly. Participants completing the pre-assessment at 9 AM have not yet received "
    "the vocabulary, framework, or psychological permission to engage with menstrual health "
    "directly.\n\n"
    "After Module 5, when participants complete the post-assessment at 4:45 PM, they read "
    "'appears unwell or in physical discomfort' with the full context of what this covers. The "
    "shift in understanding IS the learning moment. Facilitators should not explain this "
    "sequencing to participants; it works because it is implicit.\n\n"
    "If a participant asks during the pre-assessment what 'health issues' refers to: "
    "\"Any situation where a team member is not feeling their best physically. Rate based on "
    "what you do when that comes up.\""
)

TEAM_INCLUSION_NOTE = (
    "Items 19-21 address task allocation and recognition biases. Item 19 is the highest-stakes "
    "item in the instrument from an ego-threat perspective: 'I assign high-visibility or "
    "challenging tasks to women on my team, not only routine or administrative work.'\n\n"
    "Many participants genuinely do NOT do this — and genuinely believe they ARE being inclusive. "
    "The item will produce some of the lowest pre-scores in the instrument for this population. "
    "This is functioning as designed. Do not soften this item.\n\n"
    "If facilitators feel the item is too direct: the key phrase is 'not only' — this "
    "acknowledges that routine tasks are fine, the issue is whether women also get access to "
    "challenging work."
)

GENDER_HIERARCHY_NOTES = [
    "Distribute the form with minimal preamble. Do not explain why the instrument exists. "
    "Read the script above — nothing more.",
    "Do not make eye contact with specific participants during form completion. Look at the "
    "room generally, then look at your notes.",
    "Do not collect forms under any circumstances — not 'just to check if everyone finished,' "
    "not 'to make sure you have your copy.' Never hold a participant's form. This physical "
    "non-collection is the most powerful privacy signal available.",
    "Do not comment on the room's body language during completion. Let them work.",
]


def build_section7(doc):
    add_heading(doc, "Section 7: Cultural Context Notes", level=1)

    add_heading(doc, "Health Sensitivity Domain (Items 16-18)", level=2)
    add_body_para(doc, HEALTH_SENSITIVITY_NOTE)

    doc.add_paragraph()

    add_heading(doc, "Team Inclusion Domain (Items 19-21)", level=2)
    add_body_para(doc, TEAM_INCLUSION_NOTE)

    doc.add_paragraph()

    add_heading(doc, "Facilitator Gender and Hierarchy", level=2)
    add_body_para(
        doc,
        "The facilitator (likely female, likely younger than some participants) distributes "
        "a self-assessment about behavior toward women. This is a loaded dynamic. Mitigations:"
    )
    for i, note in enumerate(GENDER_HIERARCHY_NOTES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. {note}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    print("PASS: Section 7 Cultural Notes")


# ---------------------------------------------------------------------------
# SECTION 8: Form Disposition
# ---------------------------------------------------------------------------

DISPOSITION_POINTS = [
    "As participants complete the Post (PM) column, do not hover nearby. Remain at the front "
    "of the room or step outside.",
    "When time is called, say 'Please keep your form' — do not reach out your hand, do not "
    "hold a collection box, do not ask if anyone wants to hand theirs in.",
    "If the client has requested you to collect forms, decline. Explain that the privacy "
    "guarantee in the script is the foundation of honest self-report data, and collecting "
    "forms would undermine every administration of this instrument.",
    "The show-of-hands protocol (Section 5) provides the room-level aggregate data the "
    "facilitator needs. Individual forms are not required for facilitation.",
]

DISPOSITION_CLOSING = (
    "All forms go home with participants immediately after the post-assessment."
)


def build_section8(doc):
    add_heading(doc, "Section 8: Form Disposition", level=1)

    p_bold = doc.add_paragraph()
    run_b = p_bold.add_run(
        "The facilitator does NOT collect participant forms under any circumstances."
    )
    run_b.bold = True
    run_b.font.size = Pt(11)
    run_b.font.name = "Calibri"

    doc.add_paragraph()

    for i, point in enumerate(DISPOSITION_POINTS, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. {point}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()

    p_closing = doc.add_paragraph()
    run_c = p_closing.add_run(DISPOSITION_CLOSING)
    run_c.bold = True
    run_c.font.size = Pt(11)
    run_c.font.name = "Calibri"

    print("PASS: Section 8 Form Disposition")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    try:
        doc, section = create_document()

        # Page header on every page
        add_page_header(section)

        # Cover section
        build_cover(doc)

        # Section 1: Item-Domain Map
        build_section1(doc)
        doc.add_page_break()

        # Section 2: Reverse-Scored Items
        build_section2(doc)
        doc.add_page_break()

        # Section 3: Pre-Assessment Script
        build_section3(doc)
        doc.add_page_break()

        # Section 4: Post-Assessment Script
        build_section4(doc)
        doc.add_page_break()

        # Section 5: Show-of-Hands Protocol
        build_section5(doc)
        doc.add_page_break()

        # Section 6: Anonymous Tear-Off Strip
        build_section6(doc)
        doc.add_page_break()

        # Section 7: Cultural Context Notes
        build_section7(doc)
        doc.add_page_break()

        # Section 8: Form Disposition
        build_section8(doc)

        # Save
        doc.save(OUTPUT_PATH)
        print(f"PASS: File saved -> {OUTPUT_PATH}")

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
