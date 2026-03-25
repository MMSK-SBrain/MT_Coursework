#!/usr/bin/env python3
"""
Generate Training Proposal DOCX for Fortune 500 HR Leaders
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('TRAINING PROPOSAL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Inclusive Management Training for Shop Floor Supervisors')
subtitle_run.bold = True
subtitle_run.font.size = Pt(16)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

tagline = doc.add_paragraph()
tagline_run = tagline.add_run('Building Respectful Leadership in Manufacturing Environments')
tagline_run.italic = True
tagline_run.font.size = Pt(12)
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# Prepared for section
prepared = doc.add_paragraph()
prepared_run = prepared.add_run(f'Prepared for: HR Leadership Team\nDate: {datetime.now().strftime("%B %d, %Y")}')
prepared_run.font.size = Pt(11)
prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================================
# EXECUTIVE SUMMARY
# ============================================================================
doc.add_heading('Executive Summary', 1)

exec_summary = [
    "Recent employee feedback has highlighted critical gaps in supervisory management practices, particularly concerning the professional treatment and inclusion of women employees on the shop floor. These issues manifest as complaints regarding rudeness, lack of basic courtesies (such as offering seating during health-related needs), and general workplace insensitivity.",
    "",
    "This training proposal presents a comprehensive, evidence-based intervention designed to address these behavioral gaps while building a foundation for sustained inclusive leadership. The program targets male supervisors and managers across all experience levels (2-15 years) who have not previously received diversity or inclusion training.",
    "",
    "Key highlights of this program:",
]

for para in exec_summary:
    p = doc.add_paragraph(para)

bullets = [
    "One-day intensive format (8 hours) designed to minimize operational disruption",
    "50/50 balance of theory and hands-on practice for immediate behavioral application",
    "Led by external behavioral scientist with PhD in Emotional Intelligence",
    "Evidence-based curriculum grounded in neuroscience, behavioral psychology, and organizational development",
    "Measurable outcomes with 30-day follow-up and complaint tracking",
    "Culturally contextualized for Indian manufacturing environments"
]

for bullet in bullets:
    doc.add_paragraph(bullet, style='List Bullet')

impact = doc.add_paragraph()
impact.add_run("Expected Impact: ").bold = True
impact.add_run("Reduction in employee complaints, improved team cohesion, enhanced supervisor effectiveness, and stronger employer brand as an inclusive workplace.")

doc.add_page_break()

# ============================================================================
# BUSINESS CASE
# ============================================================================
doc.add_heading('Business Case', 1)

doc.add_heading('The Challenge', 2)
challenges = [
    ("Employee Complaints", "Multiple reports of disrespectful behavior and lack of basic workplace courtesies toward women employees"),
    ("Legal & Compliance Risk", "Potential exposure to workplace harassment claims and regulatory non-compliance"),
    ("Talent Retention", "Risk of losing skilled women employees due to hostile or unwelcoming work environment"),
    ("Productivity Impact", "Team friction and low morale affecting operational efficiency"),
    ("Employer Brand", "Reputation risk as industry increasingly prioritizes diversity and inclusion")
]

for title, desc in challenges:
    p = doc.add_paragraph()
    p.add_run(f"• {title}: ").bold = True
    p.add_run(desc)

doc.add_heading('The Opportunity', 2)
opportunities = [
    "Transform supervisory behavior through targeted skill development",
    "Demonstrate leadership commitment to inclusive culture",
    "Enhance employee engagement and reduce turnover costs",
    "Mitigate legal and reputational risks proactively",
    "Build competitive advantage in attracting diverse talent",
    "Improve operational performance through better team dynamics"
]

for opp in opportunities:
    doc.add_paragraph(opp, style='List Bullet')

doc.add_page_break()

# ============================================================================
# PROGRAM OBJECTIVES
# ============================================================================
doc.add_heading('Program Objectives', 1)

doc.add_paragraph("Upon completion of this training, supervisors will be able to:")

objectives = [
    ("Demonstrate respectful and inclusive management practices",
     "Apply professional communication, daily courtesies, and equitable treatment toward all team members, with specific attention to creating welcoming environment for women employees"),

    ("Apply empathy and emotional intelligence",
     "Recognize and manage personal emotions, understand diverse perspectives, and respond with appropriate empathy in workplace interactions and conflict situations"),

    ("Implement appropriate workplace accommodations",
     "Understand health-related needs (including menstrual health), provide dignified accommodations, and create supportive systems for diverse employee requirements")
]

for idx, (obj, desc) in enumerate(objectives, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{idx}. {obj}")
    run.bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_heading('Success Metrics', 2)

metrics = [
    "Measurable reduction in employee complaints within 60 days",
    "Positive feedback from women employees regarding supervisor behavior",
    "100% completion of personal action plans with specific behavioral commitments",
    "Demonstrated competency in scenario-based assessments",
    "Observable behavior change validated in 30-day follow-up reviews"
]

for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

doc.add_page_break()

# ============================================================================
# TARGET AUDIENCE
# ============================================================================
doc.add_heading('Target Audience', 1)

audience_table = doc.add_table(rows=6, cols=2)
audience_table.style = 'Light Grid Accent 1'

audience_data = [
    ("Role", "Male supervisors and managers in shop floor operations"),
    ("Experience Level", "2-15 years of management experience"),
    ("Educational Background", "Engineering graduates"),
    ("Age Range", "28-50 years"),
    ("Prior Training", "No previous diversity, inclusion, or soft skills training"),
    ("Group Size", "15-25 participants per session (optimal for interactive learning)")
]

for idx, (field, value) in enumerate(audience_data):
    row = audience_table.rows[idx]
    row.cells[0].text = field
    row.cells[1].text = value
    # Bold the first column
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_heading('Why This Audience?', 2)
rationale = [
    "This cohort has significant influence over daily employee experience and team culture",
    "Engineering backgrounds may not have included leadership or interpersonal skills development",
    "Extended tenure means established behavioral patterns that require structured intervention",
    "Critical mass of this demographic makes them leverage points for cultural transformation",
    "Addressing this group proactively prevents escalation of current issues"
]

for item in rationale:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# PROGRAM STRUCTURE
# ============================================================================
doc.add_heading('Program Structure', 1)

doc.add_heading('Overview', 2)
overview_table = doc.add_table(rows=4, cols=2)
overview_table.style = 'Light Grid Accent 1'

overview_data = [
    ("Duration", "1 day (9:00 AM - 5:00 PM)"),
    ("Training Time", "6 hours 45 minutes"),
    ("Break Time", "1 hour 15 minutes (2 breaks + lunch)"),
    ("Delivery Format", "50% Theory / 50% Hands-on Practice")
]

for idx, (field, value) in enumerate(overview_data):
    row = overview_table.rows[idx]
    row.cells[0].text = field
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_heading('Module Breakdown', 2)

modules = [
    ("Module 1: Foundation - Understanding the Context", "1 hour",
     "Why inclusive management matters • Business case and legal framework • Impact on team performance • Current state assessment • Setting learning ground rules"),

    ("Module 2: Building Self-Awareness", "1 hour",
     "Pre-training self-assessment • Understanding unconscious bias (neuroscience) • Emotional intelligence framework • Recognizing personal triggers and patterns"),

    ("Module 3: Developing Empathy", "1 hour",
     "The science of empathy • Perspective-taking and active listening • Understanding women's shop floor experiences • Empathy vs sympathy • Empathy mapping exercises • Practical application"),

    ("Module 4: Practical Communication & Daily Interactions", "1.5 hours",
     "Professional greetings and daily courtesies • Active listening techniques • Giving constructive feedback respectfully • Handling conflicts constructively • Role-play practice • Do's and Don'ts demonstrations"),

    ("Module 5: Health Sensitivity & Workplace Courtesies", "1 hour",
     "Understanding menstrual health and workplace needs • Providing appropriate accommodations (seating, breaks, privacy) • Professional responses to sensitive requests • Creating team norms for supportive environment • Scenario analysis"),

    ("Module 6: Application & Integration", "1 hour",
     "Case study analysis (Indian workplace contexts) • Complex scenario practice • Peer learning and group presentations • Scenario-based assessment • Problem-solving activities"),

    ("Module 7: Commitment & Action Planning", "15 minutes",
     "Personal action plan creation • Public commitment to behavior change • Post-training self-assessment • 30-day accountability structure • Closing")
]

for title, duration, content in modules:
    p = doc.add_paragraph()
    run1 = p.add_run(title)
    run1.bold = True
    run2 = p.add_run(f" ({duration})")
    run2.italic = True
    doc.add_paragraph(content, style='List Bullet')

doc.add_page_break()

# ============================================================================
# PEDAGOGY & METHODOLOGY
# ============================================================================
doc.add_heading('Pedagogy & Methodology', 1)

doc.add_heading('Instructional Approach', 2)
doc.add_paragraph("This program employs evidence-based adult learning principles and behavioral change methodologies:")

pedagogy = [
    ("Experiential Learning", "Participants learn by doing through role-plays, simulations, and real-world scenario practice, ensuring immediate skill application"),

    ("Behavioral Science Foundation", "Curriculum grounded in neuroscience (empathy), psychology (unconscious bias), and behavioral economics (habit formation)"),

    ("Backward Design", "Program designed from desired outcomes backward, ensuring every activity directly contributes to measurable behavioral change"),

    ("Psychologically Safe Environment", "Female facilitator with expertise in emotional intelligence creates non-judgmental space for vulnerable learning"),

    ("Contextual Relevance", "All examples, case studies, and scenarios localized to Indian manufacturing shop floor contexts"),

    ("Peer Learning", "Small group activities and discussions leverage collective wisdom and build peer accountability"),

    ("Spaced Practice", "Multiple practice opportunities throughout the day with increasing complexity, reinforcing skill development")
]

for title, desc in pedagogy:
    p = doc.add_paragraph()
    p.add_run(f"{title}: ").bold = True
    p.add_run(desc)

doc.add_heading('Training Methods', 2)

methods = [
    "Interactive lectures with multimedia content",
    "Pre/post self-assessments for baseline and progress measurement",
    "Video testimonials from women employees (anonymous)",
    "Role-play exercises with structured feedback",
    "Small group discussions and problem-solving",
    "Case study analysis using real workplace scenarios",
    "Reflective exercises (trigger mapping, empathy mapping)",
    "Scenario-based assessments",
    "Personal action planning with accountability mechanisms"
]

for method in methods:
    doc.add_paragraph(method, style='List Bullet')

doc.add_heading('Materials Provided', 2)

materials = [
    "Comprehensive participant workbook (printed)",
    "Self-assessment tools (pre/post)",
    "Case study materials",
    "Personal action plan templates",
    "Quick reference cards for daily practice",
    "Video content (empathy, unconscious bias, workplace inclusion)",
    "Scenario practice guidelines"
]

for material in materials:
    doc.add_paragraph(material, style='List Bullet')

doc.add_page_break()

# ============================================================================
# FACILITATOR PROFILE
# ============================================================================
doc.add_heading('Facilitator Profile', 1)

doc.add_heading('Qualifications', 2)

qual_points = [
    "PhD in Emotional Intelligence from recognized institution",
    "Behavioral scientist with expertise in workplace psychology",
    "Extensive experience in organizational behavior change",
    "Deep understanding of gender dynamics in workplace settings",
    "Track record of successful inclusive leadership training",
    "Cultural competence in Indian corporate environments"
]

for qual in qual_points:
    doc.add_paragraph(qual, style='List Bullet')

doc.add_heading('Why External Female Facilitator?', 2)

rationale_points = [
    ("Credibility & Expertise", "PhD-level expertise in emotional intelligence brings scientific rigor and credibility to behavioral change concepts"),

    ("Gender Perspective", "Female facilitator provides authentic perspective on women's workplace experiences, enhancing learning authenticity"),

    ("Neutral Third Party", "External consultant creates psychological safety, allowing participants to be vulnerable without internal political concerns"),

    ("Modeling Inclusive Behavior", "Participants experience respectful professional interaction with accomplished female leader, reinforcing training concepts"),

    ("Behavioral Science Authority", "Academic credentials establish authority on empathy, bias, and emotional intelligence topics")
]

for title, desc in rationale_points:
    p = doc.add_paragraph()
    p.add_run(f"{title}: ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================================
# ASSESSMENT & ACCOUNTABILITY
# ============================================================================
doc.add_heading('Assessment & Accountability', 1)

doc.add_heading('During Training', 2)

during = [
    ("Pre-Training Self-Assessment", "Baseline measurement of current attitudes, awareness, and behaviors"),
    ("Active Participation", "Engagement in all activities, discussions, and practice sessions"),
    ("Scenario-Based Assessment", "Written/verbal responses to workplace situations demonstrating learned concepts"),
    ("Personal Action Plan", "Creation of specific, measurable behavioral commitments"),
    ("Post-Training Self-Assessment", "Measurement of learning gains and self-awareness development")
]

for title, desc in during:
    p = doc.add_paragraph()
    p.add_run(f"• {title}: ").bold = True
    p.add_run(desc)

doc.add_heading('Post-Training Follow-Up', 2)

followup = [
    ("30-Day Individual Review", "One-on-one check-in with each supervisor to assess implementation of action plan and identify barriers"),

    ("Complaint Tracking", "Monitor reduction in employee complaints over 30, 60, 90-day periods"),

    ("Employee Feedback Collection", "Anonymous survey of women employees regarding observed supervisor behavior changes"),

    ("Management Observation", "Line managers observe and document behavioral improvements in daily operations"),

    ("Refresher Session (Optional)", "60-day group session to share successes, address challenges, and reinforce concepts")
]

for title, desc in followup:
    p = doc.add_paragraph()
    p.add_run(f"• {title}: ").bold = True
    p.add_run(desc)

doc.add_heading('Success Criteria', 2)

doc.add_paragraph("A supervisor is considered successful upon:")

success = [
    "Full attendance and active participation in all training modules",
    "Completion of personal action plan with minimum 3 specific behavioral commitments",
    "Demonstrated understanding in scenario-based assessments",
    "Observable behavior change validated in 30-day review",
    "Positive feedback from direct reports and peers"
]

for item in success:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# EXPECTED OUTCOMES & ROI
# ============================================================================
doc.add_heading('Expected Outcomes & Return on Investment', 1)

doc.add_heading('Immediate Outcomes (0-30 days)', 2)

immediate = [
    "Increased awareness of unconscious bias and its workplace impact",
    "Improved daily interactions and basic courtesies toward all employees",
    "Better understanding of health-related needs and appropriate accommodations",
    "Enhanced empathy and emotional intelligence in conflict situations",
    "Individual accountability through personal action plans"
]

for item in immediate:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Short-Term Outcomes (30-90 days)', 2)

short_term = [
    "Measurable reduction in employee complaints (target: 50-70% reduction)",
    "Positive feedback from women employees regarding supervisor behavior",
    "Improved team morale and cohesion on shop floor",
    "Enhanced supervisor confidence in handling diverse team dynamics",
    "Visible culture shift toward greater inclusion"
]

for item in short_term:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Long-Term Outcomes (90+ days)', 2)

long_term = [
    "Sustained behavioral change embedded in daily management practices",
    "Reduced employee turnover, particularly among women employees",
    "Stronger employer brand as inclusive workplace",
    "Improved employee engagement scores",
    "Enhanced organizational resilience and adaptability",
    "Foundation for broader diversity and inclusion initiatives",
    "Reduced legal and compliance risk",
    "Better operational performance through stronger team dynamics"
]

for item in long_term:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Return on Investment', 2)

roi_text = doc.add_paragraph()
roi_text.add_run("This training delivers ROI through multiple channels:")

roi = [
    ("Risk Mitigation", "Reducing legal exposure, compliance violations, and reputational damage (potential cost savings: significant)"),

    ("Retention", "Decreasing turnover costs associated with replacing skilled employees (industry average: 50-200% of annual salary)"),

    ("Productivity", "Improving team performance through better collaboration and reduced conflict (estimated 10-15% productivity gain)"),

    ("Engagement", "Enhancing employee engagement, which correlates with 21% higher profitability (Gallup)"),

    ("Employer Brand", "Strengthening ability to attract top talent in competitive manufacturing sector"),

    ("Cultural Foundation", "Building platform for sustained organizational development and competitive advantage")
]

for title, desc in roi:
    p = doc.add_paragraph()
    p.add_run(f"{title}: ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================================
# LOGISTICS & REQUIREMENTS
# ============================================================================
doc.add_heading('Logistics & Requirements', 1)

doc.add_heading('Facility Requirements', 2)

facility = [
    ("Room Setup", "U-shape seating configuration for 15-25 participants"),
    ("Equipment", "Projector/screen, sound system, flipcharts, markers"),
    ("Space", "Sufficient room for small group breakouts and movement during activities"),
    ("Accessibility", "Comfortable seating, adequate lighting, climate control")
]

for title, desc in facility:
    p = doc.add_paragraph()
    p.add_run(f"• {title}: ").bold = True
    p.add_run(desc)

doc.add_heading('Participant Requirements', 2)

participant = [
    "Attendance for full 8-hour program (9:00 AM - 5:00 PM)",
    "Willingness to participate actively in discussions and activities",
    "Commitment to complete pre/post assessments",
    "Openness to self-reflection and behavior change",
    "Participation in 30-day follow-up review"
]

for item in participant:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Organizational Support Needed', 2)

org_support = [
    "Senior leadership endorsement and attendance at opening/closing sessions",
    "HR support for logistics coordination and follow-up tracking",
    "Manager commitment to observe and reinforce behavior changes post-training",
    "Anonymous employee feedback mechanism for measuring impact",
    "Protection from operational disruptions during training day",
    "Commitment to 30-day individual review process"
]

for item in org_support:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Language & Cultural Considerations', 2)

doc.add_paragraph("Training will be delivered in English with all content contextualized for Indian manufacturing environments. Examples, case studies, and scenarios will reflect shop floor realities and cultural norms to ensure maximum relevance and application.")

doc.add_page_break()

# ============================================================================
# IMPLEMENTATION TIMELINE
# ============================================================================
doc.add_heading('Implementation Timeline', 1)

timeline_table = doc.add_table(rows=8, cols=3)
timeline_table.style = 'Light Grid Accent 1'

timeline_table.rows[0].cells[0].text = "Phase"
timeline_table.rows[0].cells[1].text = "Timeline"
timeline_table.rows[0].cells[2].text = "Key Activities"

timeline_data = [
    ("Kick-off & Planning", "Week 1", "Finalize participant list, confirm logistics, customize case studies to organizational context"),
    ("Pre-Training Communication", "Week 2", "Send pre-work materials, conduct pre-training self-assessments, set expectations"),
    ("Training Delivery", "Week 3", "Full-day intensive training session"),
    ("Immediate Follow-Up", "Week 4", "Collect post-training feedback, compile action plans, share initial results with HR"),
    ("30-Day Review", "Week 7", "Individual check-ins with each supervisor, track complaint data, gather employee feedback"),
    ("60-Day Assessment", "Week 11", "Measure sustained behavior change, identify additional support needs, optional refresher session"),
    ("90-Day Impact Report", "Week 15", "Comprehensive impact analysis, ROI calculation, recommendations for next steps")
]

for idx, (phase, timeline, activities) in enumerate(timeline_data, 1):
    timeline_table.rows[idx].cells[0].text = phase
    timeline_table.rows[idx].cells[1].text = timeline
    timeline_table.rows[idx].cells[2].text = activities

# Make header row bold
for cell in timeline_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ============================================================================
# NEXT STEPS
# ============================================================================
doc.add_heading('Next Steps', 1)

doc.add_paragraph("To move forward with this transformative training program:")

next_steps = [
    ("Review & Approval", "HR leadership reviews this proposal and provides feedback or approval"),

    ("Customization Meeting", "30-minute consultation to discuss organizational specifics, customize case studies, and finalize logistics"),

    ("Participant Selection", "Identify and confirm supervisors for first cohort (15-25 participants)"),

    ("Logistics Coordination", "Confirm training date, venue, equipment, and materials"),

    ("Pre-Training Communication", "Distribute pre-work and set participant expectations"),

    ("Training Delivery", "Execute full-day program"),

    ("Follow-Up & Impact Measurement", "Implement 30/60/90-day assessment framework")
]

for idx, (step, desc) in enumerate(next_steps, 1):
    p = doc.add_paragraph()
    run1 = p.add_run(f"{idx}. {step}: ")
    run1.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()

contact = doc.add_paragraph()
contact.add_run("For questions, additional information, or to schedule the customization meeting, please contact:").bold = True

doc.add_paragraph("[Your Contact Information]")
doc.add_paragraph("[Email Address]")
doc.add_paragraph("[Phone Number]")

doc.add_page_break()

# ============================================================================
# APPENDIX
# ============================================================================
doc.add_heading('Appendix: Detailed Learning Outcomes', 1)

doc.add_heading('Course-Level Outcomes', 2)

co_table = doc.add_table(rows=4, cols=2)
co_table.style = 'Light List Accent 1'

co_table.rows[0].cells[0].text = "Code"
co_table.rows[0].cells[1].text = "Description"
co_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
co_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

co_data = [
    ("IMSF-CO-1", "Demonstrate respectful and inclusive management practices toward women employees in shop floor environments"),
    ("IMSF-CO-2", "Apply empathy and emotional intelligence in daily workplace interactions and conflict situations"),
    ("IMSF-CO-3", "Implement appropriate workplace accommodations and support systems for diverse employee needs")
]

for idx, (code, desc) in enumerate(co_data, 1):
    co_table.rows[idx].cells[0].text = code
    co_table.rows[idx].cells[1].text = desc

doc.add_paragraph()

doc.add_heading('Module-Level Outcomes', 2)

mo_list = [
    "Understand the rationale for inclusive management practices",
    "Recognize personal biases and emotional patterns",
    "Apply empathy in workplace interactions",
    "Conduct respectful daily workplace interactions",
    "Provide constructive feedback professionally",
    "Provide appropriate health-related accommodations",
    "Analyze and resolve complex workplace situations",
    "Commit to ongoing inclusive management development"
]

for idx, mo in enumerate(mo_list, 1):
    doc.add_paragraph(f"{idx}. {mo}")

doc.add_paragraph()

doc.add_heading('Session-Level Outcomes', 2)
doc.add_paragraph("20 detailed session outcomes covering:")

so_categories = [
    "Understanding business case and legal frameworks",
    "Recognizing unconscious bias and triggers",
    "Developing empathy and perspective-taking skills",
    "Practicing professional communication",
    "Understanding and accommodating health needs",
    "Analyzing complex scenarios",
    "Creating actionable commitments"
]

for cat in so_categories:
    doc.add_paragraph(cat, style='List Bullet')

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_file = "/Volumes/Dev/Course_Generator/courses/inclusive-management-shop-floor/Training_Proposal_IMSF.docx"
doc.save(output_file)
print(f"Training proposal created successfully: {output_file}")
print("Document is ready for Fortune 500 HR leadership review.")
