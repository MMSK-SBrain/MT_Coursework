"""
generate_participant_form.py
Generates the participant self-assessment form for:
Inclusive Management for Shop Floor Supervisors (IMSF)

Output: assessments/pre-post-self-assessment-participant.docx

Form title: "My Management Practices"
All 21 behavioral items, 3-column Pre/Post table, self-reflection section,
domain-averaging grid, and anonymous tear-off strip.
"""

import os
import sys

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSESSMENTS_DIR = os.path.join(SCRIPT_DIR, 'assessments')
OUTPUT_PATH = os.path.join(ASSESSMENTS_DIR, 'pre-post-self-assessment-participant.docx')

# ---------------------------------------------------------------------------
# Item data — 21 behavioral items in 7 domains of 3 items each.
# Reverse-scored items (5, 11, 18) are presented identically to all others.
# No reverse-score indication appears anywhere on the participant form.
# ---------------------------------------------------------------------------
ITEMS = [
    # Domain 1: Foundation
    (1,  "I treat every person on my team with the same level of courtesy and professionalism, regardless of their role or background."),
    (2,  "When I make a decision that affects my team, I consider how it will impact different team members, not just the majority."),
    (3,  "I hold myself to the same behavioral standards I expect from others on my team."),
    # Domain 2: Daily Interactions
    (4,  "When a woman on my team speaks during a group discussion, I give her the same attention and space I give male colleagues."),
    (5,  "When I need someone to take on an extra task quickly, I turn to male team members before considering women on my team."),
    (6,  "I greet all team members at the start of the shift, including those I interact with less frequently."),
    # Domain 3: Self-Awareness
    (7,  "Before reacting to a situation that frustrates me, I pause to consider how my response might affect the people around me."),
    (8,  "I notice when my tone or manner changes depending on who I am speaking to -- and I reflect on why."),
    (9,  "I regularly consider whether my assumptions about team members are based on evidence or on habit."),
    # Domain 4: Empathy
    (10, "When a team member tells me about a problem, I listen fully before offering my opinion or a solution."),
    (11, "I find it easier to discuss personal or work difficulties with male team members than with women on my team."),
    (12, "When a woman on my team seems upset or disengaged, I check in with her privately rather than assuming it will pass."),
    # Domain 5: Feedback & Conflict
    (13, "When I give corrective feedback, I focus on the specific behavior and its impact rather than the person's character."),
    (14, "I give women on my team the same quality of developmental feedback I give to my male team members."),
    (15, "When there is a conflict on my team, I hear from all parties involved before forming my view of what happened."),
    # Domain 6: Health Sensitivity
    (16, "If a team member appears unwell or in physical discomfort, I privately offer to adjust their tasks or give them a break."),
    (17, "When a team member requests time off or a schedule adjustment for a health reason, I respond promptly without asking for personal details."),
    (18, "I hesitate to make task adjustments for a team member's health issues because I worry it will seem like preferential treatment."),
    # Domain 7: Team Inclusion
    (19, "I assign high-visibility or challenging tasks to women on my team, not only routine or administrative work."),
    (20, "When I plan team activities or recognize team achievements, I make sure women on my team are included and acknowledged equally."),
    (21, "I actively create conditions where every team member -- regardless of gender -- can contribute their best work."),
]

# Domain boundaries: add a separator row AFTER these item numbers (not after 21)
DOMAIN_BOUNDARIES = {3, 6, 9, 12, 15, 18}

SCALE_NUMBERS = "1   2   3   4   5"
POST_SHADE_HEX = "DCDCDC"   # Light grey for Post (PM) column
SEP_SHADE_HEX  = "F0F0F0"   # Very light grey for domain separator rows


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def apply_shading(cell, fill_hex: str):
    """Apply solid background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_row_height(row, height_cm: float, rule=WD_ROW_HEIGHT_RULE.AT_LEAST):
    """Set the minimum (or exact) height for a table row."""
    row.height = Cm(height_cm)
    row.height_rule = rule


def para_font(paragraph, size_pt: float, bold: bool = False, italic: bool = False,
              alignment=None):
    """Apply font settings and optional alignment to all runs in a paragraph."""
    if alignment is not None:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = 'Calibri'


def set_cell_para(cell, text: str, size_pt: float, bold: bool = False,
                  italic: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear a cell, set its text, and apply formatting."""
    # Clear existing paragraphs
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'


def add_separator_row(table):
    """Add a thin grey separator row (domain break) spanning all 3 columns."""
    row = table.add_row()
    set_row_height(row, 0.3, rule=WD_ROW_HEIGHT_RULE.EXACTLY)
    for cell in row.cells:
        apply_shading(cell, SEP_SHADE_HEX)
    # Merge all 3 cells into one clean separator
    row.cells[0].merge(row.cells[2])
    # Ensure the merged cell has no text
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.text = ''


def add_item_row(table, number: int, text: str):
    """Add one item row to the main table."""
    row = table.add_row()
    set_row_height(row, 0.85, rule=WD_ROW_HEIGHT_RULE.AT_LEAST)

    # Column 0: item text
    cell0 = row.cells[0]
    set_cell_para(cell0, f"{number}. {text}", size_pt=11,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Column 1: Pre (AM) scale
    cell1 = row.cells[1]
    set_cell_para(cell1, SCALE_NUMBERS, size_pt=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Column 2: Post (PM) scale — grey shaded
    cell2 = row.cells[2]
    set_cell_para(cell2, SCALE_NUMBERS, size_pt=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    apply_shading(cell2, POST_SHADE_HEX)


def add_paragraph(doc, text: str, size_pt: float = 11, bold: bool = False,
                  italic: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt: float = 0,
                  space_after_pt: float = 0) -> None:
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'


def set_para_spacing(doc):
    """Remove default paragraph spacing from document normal style."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Main generation
# ---------------------------------------------------------------------------

def generate():
    os.makedirs(ASSESSMENTS_DIR, exist_ok=True)
    doc = Document()

    # Remove default spacing from Normal style
    set_para_spacing(doc)

    # -----------------------------------------------------------------------
    # 1. Page setup: A4, 1.5cm margins all sides
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # -----------------------------------------------------------------------
    # 2. Header block
    # -----------------------------------------------------------------------
    # Title
    add_paragraph(doc, "My Management Practices",
                  size_pt=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=0, space_after_pt=2)

    # Subtitle
    add_paragraph(doc, "A personal reflection tool -- for your eyes only",
                  size_pt=11, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=0, space_after_pt=2)

    # Privacy statement
    add_paragraph(doc,
                  "This form is not collected by the facilitator. It is yours to keep.",
                  size_pt=10, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=0, space_after_pt=4)

    # Name / Date line
    add_paragraph(doc,
                  "Name (optional): ________________    Date: ________________",
                  size_pt=10,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=0, space_after_pt=4)

    print("PASS: Header block")

    # -----------------------------------------------------------------------
    # 3. Scale legend
    # -----------------------------------------------------------------------
    add_paragraph(doc,
                  "1 = Almost Never  |  2 = Rarely  |  3 = Sometimes  |  4 = Often  |  5 = Almost Always",
                  size_pt=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=2, space_after_pt=3)

    # -----------------------------------------------------------------------
    # 4. Item table
    # -----------------------------------------------------------------------
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths (18cm usable: 65% / 17.5% / 17.5%)
    table.columns[0].width = Cm(11.7)
    table.columns[1].width = Cm(3.15)
    table.columns[2].width = Cm(3.15)

    # Header row
    hdr = table.rows[0]

    # Header cell 0: "Statement"
    hdr_c0 = hdr.cells[0]
    set_cell_para(hdr_c0, "Statement", size_pt=11, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Header cell 1: "Pre (AM)" + scale
    hdr_c1 = hdr.cells[1]
    for para in hdr_c1.paragraphs:
        for run in para.runs:
            run.text = ''
    p1 = hdr_c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1a = p1.add_run("Pre (AM)")
    r1a.font.bold = True
    r1a.font.size = Pt(10)
    r1a.font.name = 'Calibri'
    r1a.add_break()
    r1b = p1.add_run(SCALE_NUMBERS)
    r1b.font.size = Pt(10)
    r1b.font.name = 'Calibri'

    # Header cell 2: "Post (PM)" + scale — grey shaded
    hdr_c2 = hdr.cells[2]
    for para in hdr_c2.paragraphs:
        for run in para.runs:
            run.text = ''
    p2 = hdr_c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2a = p2.add_run("Post (PM)")
    r2a.font.bold = True
    r2a.font.size = Pt(10)
    r2a.font.name = 'Calibri'
    r2a.add_break()
    r2b = p2.add_run(SCALE_NUMBERS)
    r2b.font.size = Pt(10)
    r2b.font.name = 'Calibri'
    apply_shading(hdr_c2, POST_SHADE_HEX)

    # Item rows + domain separator rows
    separator_count = 0
    for num, text in ITEMS:
        add_item_row(table, num, text)
        if num in DOMAIN_BOUNDARIES:
            add_separator_row(table)
            separator_count += 1

    print(f"PASS: Item table (21 items)")
    print(f"PASS: Separator rows ({separator_count})")

    # -----------------------------------------------------------------------
    # 5. Self-reflection section
    # -----------------------------------------------------------------------
    add_paragraph(doc, "_" * 72,
                  size_pt=10,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=6, space_after_pt=2)

    add_paragraph(doc,
                  "SELF-REFLECTION -- Complete this after filling in the Post (PM) column",
                  size_pt=11, bold=True,
                  space_before_pt=0, space_after_pt=2)

    add_paragraph(doc,
                  "Look at both columns. Where your rating changed, reflect on what shifted.",
                  size_pt=10, italic=True,
                  space_before_pt=0, space_after_pt=4)

    reflection_questions = [
        "Which group of items showed your biggest change from morning to afternoon?",
        "Which single item surprised you the most -- either this morning or now?",
        "When you return to your team, what ONE behavior will you focus on first?",
    ]
    blank_line = "_" * 67

    for i, question in enumerate(reflection_questions, 1):
        add_paragraph(doc, f"{i}. {question}",
                      size_pt=10,
                      space_before_pt=2, space_after_pt=1)
        add_paragraph(doc, blank_line,
                      size_pt=10,
                      space_before_pt=0, space_after_pt=3)

    print("PASS: Self-reflection")

    # -----------------------------------------------------------------------
    # 6. Domain-averaging grid
    # -----------------------------------------------------------------------
    add_paragraph(doc,
                  "OPTIONAL: Calculate your averages (add the three ratings, divide by 3)",
                  size_pt=9, bold=True,
                  space_before_pt=4, space_after_pt=2)

    avg_table = doc.add_table(rows=1, cols=5)
    avg_table.style = 'Table Grid'
    avg_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row for averaging grid
    avg_headers = ["Group", "Items", "Pre avg", "Post avg", "Change"]
    for i, hdr_text in enumerate(avg_headers):
        cell = avg_table.rows[0].cells[i]
        set_cell_para(cell, hdr_text, size_pt=9, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows: 7 groups
    avg_data = [
        ("Group 1", "1, 2, 3",    "/3", "/3", ""),
        ("Group 2", "4, 5, 6",    "/3", "/3", ""),
        ("Group 3", "7, 8, 9",    "/3", "/3", ""),
        ("Group 4", "10, 11, 12", "/3", "/3", ""),
        ("Group 5", "13, 14, 15", "/3", "/3", ""),
        ("Group 6", "16, 17, 18", "/3", "/3", ""),
        ("Group 7", "19, 20, 21", "/3", "/3", ""),
        ("Total",   "All items",  "__ /105", "__ /105", ""),
    ]

    for row_data in avg_data:
        row = avg_table.add_row()
        for i, cell_text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if i < 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_para(row.cells[i], cell_text, size_pt=9, alignment=align)

    print("PASS: Domain-averaging grid")

    # -----------------------------------------------------------------------
    # 7. Anonymous tear-off strip
    # -----------------------------------------------------------------------
    add_paragraph(doc,
                  "- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -",
                  size_pt=9,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=6, space_after_pt=2)

    add_paragraph(doc,
                  "OPTIONAL: Anonymous Data Strip -- Tear off and drop in the box at the front.",
                  size_pt=9, bold=True,
                  space_before_pt=0, space_after_pt=1)

    add_paragraph(doc, "Do not write your name.",
                  size_pt=9,
                  space_before_pt=0, space_after_pt=3)

    strip_lines = [
        "My Pre (AM) total score (add all 21 Pre ratings):   _______  / 105",
        "My Post (PM) total score (add all 21 Post ratings): _______  / 105",
        "My biggest single-item shift (how many points?):    _______  points",
    ]
    for line in strip_lines:
        add_paragraph(doc, line, size_pt=9,
                      space_before_pt=0, space_after_pt=2)

    print("PASS: Tear-off strip")

    # -----------------------------------------------------------------------
    # 8. Save document
    # -----------------------------------------------------------------------
    doc.save(OUTPUT_PATH)
    print(f"PASS: File saved -> {OUTPUT_PATH}")


if __name__ == '__main__':
    try:
        generate()
    except Exception as e:
        print(f"FAIL: {e}")
        sys.exit(1)
