#!/usr/bin/env python3
"""
Participant Workbook Generator
Preparatory Training for Managers: Building Respectful Leadership in Manufacturing Environments
Generates a complete participant workbook as a .docx file.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import copy

# ─────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────
NAVY        = RGBColor(0x1F, 0x38, 0x64)   # section headers
BLUE        = RGBColor(0x2E, 0x74, 0xB5)   # sub-headers / accents
LIGHT_BLUE  = "D6E4F7"                      # activity box fill (hex, no #)
CREAM       = "FFF8E7"                      # reflection box fill
GREY_FILL   = "F2F2F2"                      # case study / table header fill
DARK_GREY   = "4D4D4D"                      # table header text fill
WHITE       = "FFFFFF"
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background fill colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each value is a dict with keys: val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, props in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if props:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), props.get('val', 'single'))
            el.set(qn('w:sz'), str(props.get('sz', 6)))
            el.set(qn('w:color'), props.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_full_width(table):
    """Make table span full page width."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def add_section_header(doc, text, color=NAVY, size=14, page_break_before=False):
    """Add a bold coloured section heading paragraph."""
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_module_banner(doc, module_num, title, time_slot, color=BLUE):
    """Add a full-width coloured module banner using a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1F3864")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(8)
    r1 = p.add_run(f"MODULE {module_num}  |  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.size = Pt(11)
    r2 = p.add_run(title.upper())
    r2.bold = True
    r2.font.color.rgb = RGBColor(0xAD, 0xD8, 0xE6)
    r2.font.size = Pt(11)
    r3 = p.add_run(f"   {time_slot}")
    r3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r3.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_activity_box(doc, title, bg_color=LIGHT_BLUE, instructions=None):
    """Add a shaded activity box with optional instructions."""
    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    title_run = p.add_run(f"  {title}")
    title_run.bold = True
    title_run.font.color.rgb = NAVY
    title_run.font.size = Pt(10)
    if instructions:
        p.add_run(f"\n  {instructions}").font.size = Pt(9)
    return table, cell


def add_writing_lines(doc, count=4, label=None):
    """Add writing lines using a borderless table with bottom-border rows."""
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)
    table = doc.add_table(rows=count, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    for row in table.rows:
        cell = row.cells[0]
        set_cell_bg(cell, WHITE)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 4, 'color': 'AAAAAA'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_pull_quote(doc, text, italic=True):
    """Add an italicised pull quote centred on the page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f'"{text}"')
    r.italic = italic
    r.font.color.rgb = BLUE
    r.font.size = Pt(10)


def add_checkbox_item(doc, text, indent=0.2):
    """Add a checkbox item (□ + text)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("☐  " + text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    return p


def add_body(doc, text, size=10, space_before=2, space_after=4, indent=0, bold=False, color=None):
    """Add a plain body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    else:
        r.font.color.rgb = BLACK
    return p


def add_divider(doc, color='CCCCCC'):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)
normal.font.color.rgb = BLACK

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════

# Top colour bar
cover_table = doc.add_table(rows=1, cols=1)
set_table_full_width(cover_table)
remove_table_borders(cover_table)
cover_cell = cover_table.cell(0, 0)
set_cell_bg(cover_cell, "1F3864")
cp = cover_cell.paragraphs[0]
cp.paragraph_format.space_before = Pt(12)
cp.paragraph_format.space_after = Pt(12)
cp.paragraph_format.left_indent = Pt(8)
cp_run = cp.add_run("  PARTICIPANT WORKBOOK")
cp_run.bold = True
cp_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
cp_run.font.size = Pt(13)

doc.add_paragraph()

# Program title
t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.LEFT
t1.paragraph_format.space_before = Pt(24)
t1.paragraph_format.space_after = Pt(4)
t1r = t1.add_run("Preparatory Training for Managers")
t1r.bold = True
t1r.font.size = Pt(22)
t1r.font.color.rgb = NAVY

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.LEFT
t2.paragraph_format.space_after = Pt(32)
t2r = t2.add_run("Building Respectful Leadership in Manufacturing Environments")
t2r.italic = True
t2r.font.size = Pt(13)
t2r.font.color.rgb = BLUE

add_divider(doc, "2E74B5")

doc.add_paragraph()

# Name / Date fields
fields_table = doc.add_table(rows=3, cols=2)
set_table_full_width(fields_table)
remove_table_borders(fields_table)
field_labels = [
    ("Name:", ""),
    ("Plant / Department:", ""),
    ("Date:", datetime.now().strftime("%B %d, %Y")),
]
for i, (label, value) in enumerate(field_labels):
    lc = fields_table.cell(i, 0)
    vc = fields_table.cell(i, 1)
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(10)
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = NAVY
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(10)
    set_cell_borders(vc, bottom={'val': 'single', 'sz': 6, 'color': '2E74B5'})
    vr = vp.add_run(value)
    vr.font.size = Pt(10)
    vr.font.color.rgb = BLACK

doc.add_paragraph()
doc.add_paragraph()

# Ownership note
own = doc.add_paragraph()
own.alignment = WD_ALIGN_PARAGRAPH.LEFT
own.paragraph_format.space_before = Pt(40)
own_r = own.add_run(
    "This workbook is yours to keep.\n"
    "Your reflections, ratings, and commitments are private — nobody collects this.\n"
    "Write honestly. Use it as a mirror, a reference, and a reminder."
)
own_r.font.size = Pt(10)
own_r.italic = True
own_r.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 1 — HOW TO USE + DAY AT A GLANCE
# ═══════════════════════════════════════════════════════════════════

add_section_header(doc, "How to Use This Workbook", size=13)

add_body(doc,
    "This workbook runs alongside your training day. Each section maps to a module — "
    "the facilitator will tell you when to open to which page. Some activities are individual "
    "(quiet, private), some are in pairs or groups. All writing spaces are yours alone.",
    size=10, space_after=8
)

# Icon legend
legend_table = doc.add_table(rows=1, cols=4)
set_table_full_width(legend_table)
remove_table_borders(legend_table)
icons = [
    ("✏️  INDIVIDUAL", "Write on your own.\nYour answers are private."),
    ("👥  GROUP", "Discuss with your group.\nShare what you're comfortable with."),
    ("📋  CASE STUDY", "Read the scenario.\nAnalyse and discuss."),
    ("💡  REFERENCE", "Keep this page.\nRefer back to it on the floor."),
]
for i, (icon, desc) in enumerate(icons):
    cell = legend_table.cell(0, i)
    set_cell_bg(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(4)
    r1 = p.add_run(icon + "\n")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)

doc.add_paragraph()
add_divider(doc)

add_section_header(doc, "Today's Journey — At a Glance", size=12)

# Day timeline table
day_table = doc.add_table(rows=8, cols=3)
set_table_full_width(day_table)
remove_table_borders(day_table)

headers = ["Time", "Module", "Focus"]
for i, h in enumerate(headers):
    cell = day_table.cell(0, i)
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(4)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

rows_data = [
    ("9:00 – 10:00",   "1 — Foundation",              "Why inclusive management matters · Business case · Legal framework"),
    ("10:15 – 11:15",  "2 — Self-Awareness",           "Unconscious bias · Social intelligence · My triggers"),
    ("11:15 – 12:15",  "3 — Developing Empathy",       "Neuroscience of empathy · Her experience · Gen Z expectations"),
    ("12:15 – 1:00",   "LUNCH",                        "Case Study 1 on your chair — read before you return"),
    ("1:00 – 2:30",    "4 — Practical Communication",  "Greetings · Listening · Feedback · Conflict · Role-play"),
    ("2:30 – 3:30",    "5 — Health Sensitivity",       "Menstrual health · Accommodations · Response scripts · Team norms"),
    ("3:45 – 5:00",    "6 & 7 — Application & Commitment", "Case Study 2 · Grievance implications · Assessment · Action plan"),
]
for ri, (time, mod, focus) in enumerate(rows_data):
    bg = GREY_FILL if "LUNCH" in mod else WHITE
    for ci, text in enumerate([time, mod, focus]):
        cell = day_table.cell(ri + 1, ci)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 2, 'color': 'DDDDDD'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(9)
        if ci == 0:
            r.bold = True
            r.font.color.rgb = NAVY
        elif "LUNCH" in mod:
            r.bold = True
            r.font.color.rgb = BLUE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 2 — MODULE 1: BASELINE SELF-CHECK
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 1, "Foundation — Understanding the Context", "9:00 – 10:00 AM")

add_body(doc,
    "Before we begin, rate yourself honestly on each statement below. "
    "There are no right or wrong answers — this is your starting point. "
    "You will rate yourself again at the end of the day.",
    size=9, space_after=6
)

# Self-check table
sc_table = doc.add_table(rows=6, cols=3)
set_table_full_width(sc_table)
remove_table_borders(sc_table)

# Header row
headers_sc = ["Statement", "Morning\n(1–5)", "Evening\n(1–5)"]
header_bg = ["1F3864", "2E74B5", "2E74B5"]
for ci, (h, bg) in enumerate(zip(headers_sc, header_bg)):
    cell = sc_table.cell(0, ci)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(4)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

statements = [
    "I greet all team members — men and women — in the same way every morning.",
    "I know at least one accommodation I can make if a woman on my team needs a short rest for health reasons.",
    "When I give feedback to a woman employee, I do it the same way I would to a male employee.",
    "I am aware of at least one unconscious assumption I hold about women on the shop floor.",
    "I feel confident handling a complaint or discomfort raised by a woman on my team.",
]
for ri, stmt in enumerate(statements):
    bg = LIGHT_BLUE if ri % 2 == 0 else WHITE
    for ci in range(3):
        cell = sc_table.cell(ri + 1, ci)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 2, 'color': 'CCCCCC'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(4)
        text = stmt if ci == 0 else ""
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = BLACK

add_body(doc, "Scale:  1 = Not at all   ·   2 = Rarely   ·   3 = Sometimes   ·   4 = Often   ·   5 = Always",
         size=8, space_before=4, space_after=10, color=RGBColor(0x66, 0x66, 0x66))

add_divider(doc)

# ─── Reflection activity ───
add_activity_box(doc, "✏️  REFLECTION — The Manager I Want to Be",
                 bg_color=CREAM,
                 instructions="Think about a manager who made you feel genuinely valued early in your career.")
doc.add_paragraph()
add_writing_lines(doc, count=3, label="What did that manager do — specifically?")
add_writing_lines(doc, count=2, label="What's one thing you want to be known for as a manager?")
add_pull_quote(doc, "Your team members are building their picture of you — one interaction at a time.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 3 — MODULE 1: LEGAL + MISCONDUCT REFERENCE
# ═══════════════════════════════════════════════════════════════════

add_section_header(doc, "💡  Reference — Legal Framework & What Counts as Misconduct", color=BLUE, size=11)

# PoSH key points
posh_table = doc.add_table(rows=1, cols=2)
set_table_full_width(posh_table)
remove_table_borders(posh_table)
posh_left = posh_table.cell(0, 0)
posh_right = posh_table.cell(0, 1)
set_cell_bg(posh_left, LIGHT_BLUE)
set_cell_bg(posh_right, GREY_FILL)

for cell, title, bullets in [
    (posh_left, "PoSH Act 2013 — Key Points", [
        "Mandatory for all workplaces with 10+ employees",
        "Covers harassment AND hostile work environment",
        "You are responsible for your team's environment",
        "Prevention is your job. Redressal is the ICC's job.",
        "Non-compliance: criminal penalties + licence risk",
    ]),
    (posh_right, "What Counts as Misconduct", [
        "OBVIOUS: Physical contact · Sexual comments · Discriminatory remarks",
        "LESS OBVIOUS: Consistent dismissiveness · Denying a seat when needed",
        "LESS OBVIOUS: Giving women less challenging work · Talking over them",
        "LESS OBVIOUS: Making health needs a joke · Excluding from team talk",
        "TEST: 'Would I do / allow this if she were male?'",
    ]),
]:
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.left_indent = Pt(6)
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = NAVY
    for b in bullets:
        pb = cell.add_paragraph()
        pb.paragraph_format.left_indent = Pt(12)
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after = Pt(2)
        rb = pb.add_run("•  " + b)
        rb.font.size = Pt(8)
        rb.font.color.rgb = BLACK

doc.add_paragraph()
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════
# PAGE 4 — MODULE 2: OBSERVATION EXERCISE + BIAS TYPES
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 2, "Building Self-Awareness", "10:15 – 11:15 AM")

add_activity_box(doc, "✏️  OBSERVATION EXERCISE",
                 bg_color=LIGHT_BLUE,
                 instructions="For each moment, write one thing you notice and one thing you would do.")
doc.add_paragraph()

moments = [
    ("Moment 1",
     "It's 10 AM. Kavitha, one of your operators, has been standing at her station for 4 hours. "
     "She looks pale and is holding her stomach. She hasn't asked for anything."),
    ("Moment 2",
     "You're about to give feedback to Priya, a team member who made a quality error. "
     "Three other operators are standing nearby."),
    ("Moment 3",
     "In a team briefing you ask for ideas on a workflow problem. Two men speak immediately. "
     "Deepa, who usually has strong ideas, says nothing."),
]

for label, scenario in moments:
    obs_t = doc.add_table(rows=1, cols=1)
    set_table_full_width(obs_t)
    remove_table_borders(obs_t)
    cell = obs_t.cell(0, 0)
    set_cell_bg(cell, GREY_FILL)
    cp = cell.paragraphs[0]
    cp.paragraph_format.left_indent = Pt(6)
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    r1 = cp.add_run(label + ":  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = NAVY
    r2 = cp.add_run(scenario)
    r2.font.size = Pt(9)
    r2.font.color.rgb = BLACK
    add_writing_lines(doc, count=2, label="I notice:  /  I would:")

add_divider(doc)

add_section_header(doc, "💡  Reference — Common Biases on the Shop Floor", color=BLUE, size=10)

bias_table = doc.add_table(rows=5, cols=2)
set_table_full_width(bias_table)
remove_table_borders(bias_table)
biases = [
    ("AFFINITY BIAS",       "I connect more easily with people who are like me."),
    ("ATTRIBUTION BIAS",    "He made a mistake → bad day.  She made a mistake → incompetent."),
    ("HALO / HORN EFFECT",  "One good (or bad) thing colours everything I see after."),
    ("AUTHORITY BIAS",      "I trust the man's opinion more, even when skill is equal."),
    ("PROVE-IT-AGAIN BIAS", "Women must demonstrate competence repeatedly. Men are assumed competent."),
]
for ri, (name, desc) in enumerate(biases):
    bg = LIGHT_BLUE if ri % 2 == 0 else WHITE
    nc = bias_table.cell(ri, 0)
    dc = bias_table.cell(ri, 1)
    for cell, text, bold in [(nc, name, True), (dc, desc, False)]:
        set_cell_bg(cell, bg)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 2, 'color': 'DDDDDD'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.bold = bold
        r.font.color.rgb = NAVY if bold else BLACK

add_writing_lines(doc, count=2,
                  label="Which of these have I seen — in myself or around me? Be specific:")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 5 — MODULE 2: NOTICE-RESPOND LOOP + TRIGGER MAP
# ═══════════════════════════════════════════════════════════════════

add_section_header(doc, "The Notice-Respond Loop", color=BLUE, size=11)

loop_table = doc.add_table(rows=1, cols=4)
set_table_full_width(loop_table)
remove_table_borders(loop_table)
steps = [
    ("1. NOTICE", "What is happening right now?"),
    ("2. INTERPRET", "What might she need?"),
    ("3. RESPOND", "What is the right action here?"),
    ("4. REFLECT", "How did that land?"),
]
for i, (step, prompt) in enumerate(steps):
    cell = loop_table.cell(0, i)
    set_cell_bg(cell, "1F3864" if i % 2 == 0 else "2E74B5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(step + "\n")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2 = p.add_run(prompt)
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

add_writing_lines(doc, count=2,
                  label="In my daily management, I most often skip which step — and why?")

add_divider(doc)

# Trigger Map
add_activity_box(doc, "✏️  MY TRIGGER MAP — Individual, Private",
                 bg_color=CREAM,
                 instructions="Fill in all four boxes honestly. Nobody reads this. The gap between Box 2 and Box 4 is where your leadership lives.")
doc.add_paragraph()

trigger_table = doc.add_table(rows=2, cols=2)
set_table_full_width(trigger_table)
remove_table_borders(trigger_table)
boxes = [
    ("BOX 1 — A situation at work that regularly frustrates me:",
     "(Be specific — not 'when things go wrong' but what exactly)"),
    ("BOX 2 — How I typically react:",
     "(Honestly — what do you say / do / show on your face)"),
    ("BOX 3 — Who usually sees that reaction?",
     "(Your team? Specific people? Note who is watching)"),
    ("BOX 4 — What I wish I did instead:",
     "(The response you'd be proud of — describe it)"),
]
for idx, (title, hint) in enumerate(boxes):
    row = idx // 2
    col = idx % 2
    cell = trigger_table.cell(row, col)
    set_cell_bg(cell, WHITE)
    border_color = "2E74B5"
    set_cell_borders(cell,
                     top={'val': 'single', 'sz': 8, 'color': border_color},
                     bottom={'val': 'single', 'sz': 4, 'color': border_color},
                     left={'val': 'single', 'sz': 4, 'color': border_color},
                     right={'val': 'single', 'sz': 4, 'color': border_color})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Pt(6)
    r1 = p.add_run(title + "\n")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(hint + "\n\n\n\n")
    r2.font.size = Pt(8)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_pull_quote(doc, "The gap between stimulus and response is where leadership lives.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 6 — MODULE 3: EMPATHY VS SYMPATHY + KAVITHA EXERCISE
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 3, "Developing Empathy", "11:15 AM – 12:15 PM")

add_section_header(doc, "Empathy vs. Sympathy", color=BLUE, size=11)

es_table = doc.add_table(rows=5, cols=2)
set_table_full_width(es_table)
remove_table_borders(es_table)
for ci, (label, bg) in enumerate([("SYMPATHY", "D6E4F7"), ("EMPATHY", "1F3864")]):
    cell = es_table.cell(0, ci)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(6)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY if ci == 0 else RGBColor(0xFF, 0xFF, 0xFF)

rows_es = [
    ('"That\'s too bad."',                '"That sounds really hard."'),
    ('Fixes the discomfort quickly',       'Stays present with the discomfort'),
    ('"You\'ll be fine — carry on."',     '"Tell me more. I\'m listening."'),
    ('Closes the conversation',            'Opens the conversation'),
]
for ri, (sym, emp) in enumerate(rows_es):
    bg = GREY_FILL if ri % 2 == 0 else WHITE
    for ci, text in enumerate([sym, emp]):
        cell = es_table.cell(ri + 1, ci)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 2, 'color': 'CCCCCC'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(8)
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = BLACK

add_body(doc,
    "In management: empathy is not a feeling — it is a behaviour. "
    "It starts with one question: 'What is this like for you?'",
    size=9, space_before=6, space_after=10, bold=True, color=BLUE)

add_divider(doc)

# Kavitha exercise
add_activity_box(doc, "✏️  PERSPECTIVE EXERCISE — Through Kavitha's Eyes",
                 bg_color=CREAM,
                 instructions="Read the scenario below and answer the questions from Kavitha's perspective — not yours.")
doc.add_paragraph()

kavitha_t = doc.add_table(rows=1, cols=1)
set_table_full_width(kavitha_t)
remove_table_borders(kavitha_t)
kc = kavitha_t.cell(0, 0)
set_cell_bg(kc, GREY_FILL)
kp = kc.paragraphs[0]
kp.paragraph_format.left_indent = Pt(8)
kp.paragraph_format.space_before = Pt(6)
kp.paragraph_format.space_after = Pt(6)
kr = kp.add_run(
    "Kavitha has been on your line for two years. She is technically skilled. She joined a team of "
    "18 men. The physical environment — tools, workstations, PPE — was designed for the average male body. "
    "She arrives 20 minutes early every morning to find PPE that fits. The closest women's washroom is "
    "a 7-minute walk. Twice this month, a male colleague made a comment about her 'taking too long' for a break."
)
kr.font.size = Pt(9)
kr.font.color.rgb = BLACK

questions_k = [
    "What does Kavitha feel when she arrives at the plant gate each morning?",
    "What does she wish her supervisor understood about her day?",
    "What would make her feel genuinely valued and included on this team?",
    "What one thing could her manager do this week that would change her experience?",
]
for q in questions_k:
    add_writing_lines(doc, count=2, label=q)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 7 — MODULE 3: EMPATHY MAP + GEN Z
# ═══════════════════════════════════════════════════════════════════

add_activity_box(doc, "✏️  EMPATHY MAP — A 5-Minute Tool You Can Use on the Floor",
                 bg_color=LIGHT_BLUE,
                 instructions="Think of one specific woman on your team. Don't write her name. Fill in the four quadrants.")
doc.add_paragraph()

em_table = doc.add_table(rows=2, cols=2)
set_table_full_width(em_table)
remove_table_borders(em_table)
em_labels = [
    ("THINKS", "What is on her mind that she doesn't say aloud?"),
    ("FEELS",  "What emotions is she carrying into the shift?"),
    ("SAYS",   "What does she actually tell you?"),
    ("DOES",   "What behaviours do you observe?"),
]
for idx, (quadrant, prompt) in enumerate(em_labels):
    row = idx // 2
    col = idx % 2
    cell = em_table.cell(row, col)
    bg = "EBF3FB" if (row + col) % 2 == 0 else "D6E4F7"
    set_cell_bg(cell, bg)
    border_color = "2E74B5"
    set_cell_borders(cell,
                     top={'val': 'single', 'sz': 6, 'color': border_color},
                     bottom={'val': 'single', 'sz': 6, 'color': border_color},
                     left={'val': 'single', 'sz': 6, 'color': border_color},
                     right={'val': 'single', 'sz': 6, 'color': border_color})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Pt(6)
    r1 = p.add_run(quadrant + "\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(prompt + "\n\n\n\n\n")
    r2.font.size = Pt(8)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_body(doc, "THE GAP BETWEEN FEELS AND SAYS = YOUR OPPORTUNITY AS A MANAGER",
         size=9, space_before=8, space_after=4, bold=True, color=BLUE)
add_writing_lines(doc, count=2, label="What will I do differently because of this map?")

add_divider(doc)

add_section_header(doc, "💡  Gen Z on Your Floor — Key Expectations", color=BLUE, size=10)

genz_items = [
    "60% expect managers to care about their well-being — not just their output.",
    "77% want to work for a company whose values align with their own.",
    "They will leave — quickly — if respect is consistently absent.",
    "'Boundary-setting' is self-advocacy, not insubordination.",
    "They want feedback, growth, and to be taken seriously — by name, not by role.",
    "They have seen what good management looks like — online, globally. They know the difference.",
]
for item in genz_items:
    add_body(doc, "▸  " + item, size=9, space_before=2, space_after=2, indent=0.1)

add_writing_lines(doc, count=2,
                  label="A team member I think is Gen Z — and one assumption I might be making about them:")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 8 — CASE STUDY 1: "FIVE MINUTES"
# ═══════════════════════════════════════════════════════════════════

# Banner
cs1_table = doc.add_table(rows=1, cols=1)
set_table_full_width(cs1_table)
remove_table_borders(cs1_table)
cs1_cell = cs1_table.cell(0, 0)
set_cell_bg(cs1_cell, "1F3864")
cs1p = cs1_cell.paragraphs[0]
cs1p.paragraph_format.space_before = Pt(6)
cs1p.paragraph_format.space_after = Pt(6)
cs1p.paragraph_format.left_indent = Pt(8)
r1 = cs1p.add_run("📋  CASE STUDY 1  |  ")
r1.bold = True
r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
r1.font.size = Pt(11)
r2 = cs1p.add_run('"FIVE MINUTES"')
r2.bold = True
r2.font.color.rgb = RGBColor(0xAD, 0xD8, 0xE6)
r2.font.size = Pt(11)
r3 = cs1p.add_run("   Read during lunch. Discuss at 1:00 PM.")
r3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
r3.font.size = Pt(9)

doc.add_paragraph()

scenario_t = doc.add_table(rows=1, cols=1)
set_table_full_width(scenario_t)
remove_table_borders(scenario_t)
sc1 = scenario_t.cell(0, 0)
set_cell_bg(sc1, GREY_FILL)
set_cell_borders(sc1,
                 top={'val': 'single', 'sz': 8, 'color': '1F3864'},
                 bottom={'val': 'single', 'sz': 4, 'color': 'AAAAAA'},
                 left={'val': 'single', 'sz': 4, 'color': 'AAAAAA'},
                 right={'val': 'single', 'sz': 4, 'color': 'AAAAAA'})
sp1 = sc1.paragraphs[0]
sp1.paragraph_format.left_indent = Pt(8)
sp1.paragraph_format.right_indent = Pt(8)
sp1.paragraph_format.space_before = Pt(8)
sp1.paragraph_format.space_after = Pt(8)

scenario_text = (
    "SETTING:  Automotive component manufacturing plant, South India. Wednesday, 10:40 AM. Assembly Line B.\n\n"
    "RAJAN — Group Leader, 11 years' experience. Known as fair but brisk. Runs a tight line.\n"
    "MEERA — Operator, 22 years old, 8 months on the job. First job after ITI. One of two women on the line.\n\n"
    "Meera approaches Rajan quietly at his workstation. She looks pale. She speaks in a low voice:\n"
    "'Sir, I'm not feeling well. Can I sit down for five minutes? I'll make up the time.'\n\n"
    "Rajan is reviewing a quality report. There's a micro-stoppage on another machine he's tracking. "
    "He looks up briefly and says:\n"
    "'Meera, we're 40 units behind. Everyone needs to be at their station. "
    "Can you push through for another hour? Take your regular break at 11.'\n\n"
    "Meera nods and walks back to her station. She doesn't take a break at 11. "
    "She stays quiet for the rest of the shift.\n"
    "The next morning, she calls in sick. Three days later, Rajan receives a written complaint through HR."
)
sr = sp1.add_run(scenario_text)
sr.font.size = Pt(9)
sr.font.color.rgb = BLACK

doc.add_paragraph()
add_section_header(doc, "Discussion Questions", color=NAVY, size=10)

cs1_questions = [
    "What did Rajan get wrong? What — if anything — did he get right?",
    "What did Meera need in that moment — beyond just sitting down?",
    "If Rajan had applied the Gen Z context, health sensitivity, and empathy principles — what might he have said instead?",
    "Write Rajan's better response — the one that handles production pressure AND Meera's need with dignity (2–3 sentences).",
    "What would have happened to the complaint if Rajan had responded differently in that 30-second interaction?",
]
for i, q in enumerate(cs1_questions):
    add_body(doc, f"Q{i+1}.  {q}", size=9, space_before=4, space_after=2, bold=False)
    lines = 4 if i == 3 else 2
    add_writing_lines(doc, count=lines)

add_pull_quote(doc, "A 30-second response, done right, would have cost nothing and prevented everything that followed.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 9 — MODULE 4: ROLE-PLAY OBSERVER CHECKLIST + SBI TEMPLATE
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 4, "Practical Communication & Daily Interactions", "1:00 – 2:30 PM")

add_activity_box(doc, "👥  ROLE-PLAY OBSERVER CHECKLIST",
                 bg_color=LIGHT_BLUE,
                 instructions="Observer: watch carefully and tick what you see. Give 60 seconds of feedback after each round.")
doc.add_paragraph()

# Scenario A — Conflict
add_body(doc, "SCENARIO A — The Conflict Situation", size=9, bold=True, space_before=4, space_after=4, color=NAVY)
add_body(doc,
    "Rekha reports that a male colleague (Suresh) made a sarcastic comment in front of the team: "
    "'Madam, maybe this work is too technical for you.' She wants you to do something about it.",
    size=8, space_before=0, space_after=4, color=RGBColor(0x44, 0x44, 0x44))
conflict_checks = [
    "Acknowledged Rekha's concern before offering any solution",
    "Did NOT dismiss or minimise ('don't take it personally', 'he didn't mean it')",
    "Gave Rekha a specific next step — not vague reassurance",
    "Committed to address Suresh separately (not in front of Rekha or the team)",
    "Used behaviour language about Suresh — not character judgment",
]
for c in conflict_checks:
    add_checkbox_item(doc, c)

doc.add_paragraph()
add_body(doc, "SCENARIO B — The Feedback Conversation", size=9, bold=True, space_before=4, space_after=4, color=NAVY)
add_body(doc,
    "Sunita has been late to station twice this week by 5–7 minutes. "
    "Use the SBI model. Deliver after shift, in private.",
    size=8, space_before=0, space_after=4, color=RGBColor(0x44, 0x44, 0x44))
feedback_checks = [
    "Used SBI structure (Situation / Behaviour / Impact)",
    "Delivered privately — not on the open floor",
    "No generalisations ('always' / 'never' / character comments)",
    "Ended with a clear, specific, forward-looking ask",
]
for c in feedback_checks:
    add_checkbox_item(doc, c)

doc.add_paragraph()
add_body(doc, "SCENARIO C — Responding to a Concern", size=9, bold=True, space_before=4, space_after=4, color=NAVY)
add_body(doc,
    "Anjali approaches you quietly — the workstation she was moved to is causing shoulder pain. "
    "She looks uncertain about raising it.",
    size=8, space_before=0, space_after=4, color=RGBColor(0x44, 0x44, 0x44))
concern_checks = [
    "Stopped and gave full physical attention before responding",
    "Let Anjali finish without interrupting or redirecting",
    "Reflected back what was heard before responding",
    "Acknowledged the concern before moving to a solution",
]
for c in concern_checks:
    add_checkbox_item(doc, c)

add_writing_lines(doc, count=2, label="What I noticed as observer — the moment that stood out:")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 10 — SBI PRACTICE TEMPLATE + RESPONSE SCRIPTS
# ═══════════════════════════════════════════════════════════════════

add_section_header(doc, "SBI Feedback — Practice Template", color=BLUE, size=11)
add_body(doc,
    "SBI = Situation · Behaviour · Impact.  "
    "Use this model for every piece of corrective feedback. "
    "Keep it private, specific, and forward-looking.",
    size=9, space_before=2, space_after=8)

for template_num in range(1, 3):
    sbi_t = doc.add_table(rows=3, cols=2)
    set_table_full_width(sbi_t)
    remove_table_borders(sbi_t)
    sbi_labels = [
        ("S  SITUATION", '"During [specific time], on [location]..."'),
        ("B  BEHAVIOUR", '"I noticed that [observable action — not character]..."'),
        ("I  IMPACT",    '"...which resulted in [measurable outcome]."'),
    ]
    for ri, (label, prompt) in enumerate(sbi_labels):
        lc = sbi_t.cell(ri, 0)
        vc = sbi_t.cell(ri, 1)
        set_cell_bg(lc, "1F3864" if ri == 0 else ("2E74B5" if ri == 1 else LIGHT_BLUE.replace("D6","C0")))
        set_cell_bg(vc, WHITE)
        set_cell_borders(vc, bottom={'val': 'single', 'sz': 4, 'color': 'AAAAAA'})
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after = Pt(8)
        lp.paragraph_format.left_indent = Pt(4)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if ri < 2 else NAVY
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after = Pt(4)
        vp.paragraph_format.left_indent = Pt(6)
        vr = vp.add_run(prompt)
        vr.font.size = Pt(8)
        vr.italic = True
        vr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

add_divider(doc)

add_section_header(doc, "Module 5 — Response Scripts: Practice", color=BLUE, size=11)
add_body(doc,
    "Pattern:  ACKNOWLEDGE  →  ACTION or NEXT STEP  →  CLEAN CLOSE\n"
    "Short is professional. The urge to over-explain is discomfort leaking into the conversation.",
    size=9, space_before=2, space_after=8)

scripts = [
    ('"Sir, I\'m not feeling well. Can I sit down for a bit?"',
     "Of course — go use the chair near [location]. Let me know when you're back."),
    ('"Sir, I need to step out to the washroom — I can\'t wait for the 3 o\'clock break."',
     "Go ahead. I'll note it — come back when you're ready."),
    ('"Sir, the workstation height has been causing me back pain — I\'ve been managing, but it\'s getting worse."',
     "Thank you for telling me. That shouldn't be ongoing — let me check what adjustment is possible and come back to you before end of shift."),
]
for request, model in scripts:
    req_t = doc.add_table(rows=1, cols=1)
    set_table_full_width(req_t)
    remove_table_borders(req_t)
    rc = req_t.cell(0, 0)
    set_cell_bg(rc, GREY_FILL)
    rp = rc.paragraphs[0]
    rp.paragraph_format.left_indent = Pt(8)
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(4)
    rr = rp.add_run("REQUEST:  " + request)
    rr.font.size = Pt(9)
    rr.font.color.rgb = BLACK

    add_writing_lines(doc, count=2, label="My response:")

    model_t = doc.add_table(rows=1, cols=1)
    set_table_full_width(model_t)
    remove_table_borders(model_t)
    mc = model_t.cell(0, 0)
    set_cell_bg(mc, LIGHT_BLUE)
    mp = mc.paragraphs[0]
    mp.paragraph_format.left_indent = Pt(8)
    mp.paragraph_format.space_before = Pt(4)
    mp.paragraph_format.space_after = Pt(4)
    mr1 = mp.add_run("MODEL RESPONSE:  ")
    mr1.bold = True
    mr1.font.size = Pt(8)
    mr1.font.color.rgb = NAVY
    mr2 = mp.add_run(model)
    mr2.font.size = Pt(8)
    mr2.italic = True
    mr2.font.color.rgb = BLACK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 11 — MODULE 5: TEAM NORMS
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 5, "Health Sensitivity & Workplace Courtesies", "2:30 – 3:30 PM")

add_section_header(doc, "💡  Quick Reference — Accommodation Categories", color=BLUE, size=10)

acc_table = doc.add_table(rows=4, cols=3)
set_table_full_width(acc_table)
remove_table_borders(acc_table)

for ci, header in enumerate(["Category", "Default", "Better"]):
    cell = acc_table.cell(0, ci)
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(4)
    r = p.add_run(header)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

acc_rows = [
    ("SEATING",          "Ask HR if adjustments needed",            "Stool accessible. Quiet offer. No discussion."),
    ("WASHROOM ACCESS",  "Fixed break schedule only",               "Allow a brief exit on clearly difficult days. No interrogation."),
    ("PRIVACY",          "Address wherever conversation starts",    "Step away from team. Lower voice. Keep it brief."),
]
for ri, (cat, default, better) in enumerate(acc_rows):
    bg = LIGHT_BLUE if ri % 2 == 0 else WHITE
    for ci, text in enumerate([cat, default, better]):
        cell = acc_table.cell(ri + 1, ci)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 2, 'color': 'CCCCCC'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(8)
        r.bold = (ci == 0)
        r.font.color.rgb = NAVY if ci == 0 else BLACK

doc.add_paragraph()
add_divider(doc)

add_activity_box(doc, "✏️  TEAM NORMS — Building What Belongs to Us",
                 bg_color=CREAM,
                 instructions="Think about YOUR team, YOUR floor. Answer all three questions honestly.")
doc.add_paragraph()

norms_qs = [
    "One thing about how my team currently operates that I want to KEEP — something that already works well:",
    "One behaviour I want to STOP — something I've seen or done that I now recognise as disrespectful or exclusionary:",
    "One new NORM I want to establish — specific, observable, introducible next week:",
]
for q in norms_qs:
    add_writing_lines(doc, count=3, label=q)

add_section_header(doc, "Our Floor's New Norms — From the Group Discussion", color=NAVY, size=10)
add_body(doc, "Write the norms your group agreed on during the debrief:", size=9, space_after=4)
norm_table = doc.add_table(rows=5, cols=2)
set_table_full_width(norm_table)
remove_table_borders(norm_table)
for ri in range(5):
    nc = norm_table.cell(ri, 0)
    vc = norm_table.cell(ri, 1)
    set_cell_bg(nc, "1F3864")
    set_cell_bg(vc, WHITE)
    set_cell_borders(vc, bottom={'val': 'single', 'sz': 4, 'color': 'AAAAAA'})
    np = nc.paragraphs[0]
    np.paragraph_format.space_before = Pt(10)
    np.paragraph_format.space_after = Pt(10)
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(str(ri + 1))
    nr.bold = True
    nr.font.size = Pt(12)
    nr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(8)
    vp.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 12 — CASE STUDY 2: "THE REPORT"
# ═══════════════════════════════════════════════════════════════════

cs2_t = doc.add_table(rows=1, cols=1)
set_table_full_width(cs2_t)
remove_table_borders(cs2_t)
cs2_cell = cs2_t.cell(0, 0)
set_cell_bg(cs2_cell, "1F3864")
cs2p = cs2_cell.paragraphs[0]
cs2p.paragraph_format.space_before = Pt(6)
cs2p.paragraph_format.space_after = Pt(6)
cs2p.paragraph_format.left_indent = Pt(8)
r1 = cs2p.add_run("📋  CASE STUDY 2  |  ")
r1.bold = True
r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
r1.font.size = Pt(11)
r2 = cs2p.add_run('"THE REPORT"')
r2.bold = True
r2.font.color.rgb = RGBColor(0xAD, 0xD8, 0xE6)
r2.font.size = Pt(11)
r3 = cs2p.add_run("   Read in Module 6 — Place on chair before 3:45 PM.")
r3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
r3.font.size = Pt(9)

doc.add_paragraph()

sc2_table = doc.add_table(rows=1, cols=1)
set_table_full_width(sc2_table)
remove_table_borders(sc2_table)
sc2 = sc2_table.cell(0, 0)
set_cell_bg(sc2, GREY_FILL)
set_cell_borders(sc2,
                 top={'val': 'single', 'sz': 8, 'color': '1F3864'},
                 bottom={'val': 'single', 'sz': 4, 'color': 'AAAAAA'},
                 left={'val': 'single', 'sz': 4, 'color': 'AAAAAA'},
                 right={'val': 'single', 'sz': 4, 'color': 'AAAAAA'})
sp2 = sc2.paragraphs[0]
sp2.paragraph_format.left_indent = Pt(8)
sp2.paragraph_format.right_indent = Pt(8)
sp2.paragraph_format.space_before = Pt(8)
sp2.paragraph_format.space_after = Pt(8)
sr2 = sp2.add_run(
    "SETTING:  Quality Inspection Zone, automotive components plant, North India. Thursday, 2:55 PM. End of second shift.\n\n"
    "VIKRAM — Group Leader, 14 years' experience. Highly competent. Performance awards. No previous complaints.\n"
    "ANJALI — Quality Inspector, 26 years old, 14 months with the company. High-potential employee. "
    "One of three women in the inspection zone.\n"
    "FOUR MALE OPERATORS — present at the time of the incident.\n\n"
    "Vikram notices two missing entries in Anjali's inspection log for the 2:00 PM batch. "
    "He walks directly to her station and says loudly:\n"
    "'Anjali — what is this? Two components not logged? What were you doing at 2 o'clock? "
    "This is basic work. Even a new joiner doesn't make this mistake. "
    "I don't know why I have to keep checking your work when I don't check anyone else's.'\n\n"
    "Anjali replies: 'Sir, those components were sent back for rework at 1:50 — "
    "I flagged it in the physical register but hadn't updated the log yet.'\n\n"
    "Vikram responds: 'That's not a process. Update the log first, rework second. "
    "I shouldn't have to explain this to you.' He walks away.\n\n"
    "The next morning: an email arrives in HR from Anjali, marked to the ICC.\n"
    "Subject: 'Complaint — hostile and disrespectful treatment by Group Leader.'\n"
    "Two days later, Vikram is called for a preliminary conversation by the ICC Chairperson."
)
sr2.font.size = Pt(9)
sr2.font.color.rgb = BLACK

doc.add_paragraph()
add_section_header(doc, "Discussion Questions", color=NAVY, size=10)

cs2_questions = [
    "List every specific behaviour in Vikram's response that could form the basis of a complaint. "
    "Be precise — not 'he was rude,' but exactly what he said and how.",
    "Vikram believes he did nothing wrong — he was addressing a real quality lapse. Is he right? "
    "Where does legitimate performance management end and misconduct begin?",
    "Walk through the ICC process from Anjali's complaint. What happens next? "
    "Who is involved? What can the outcome be — for Anjali, for Vikram, for the organisation?",
    "Rewrite the interaction in 3–5 sentences using SBI, private feedback, and acknowledging first. "
    "Same concern — completely different approach.",
    "What is the cost to Vikram's credibility, team morale, and organisational trust — "
    "even if the ICC eventually rules in his favour?",
]
for i, q in enumerate(cs2_questions):
    add_body(doc, f"Q{i+1}.  {q}", size=9, space_before=4, space_after=2)
    lines = 5 if i in [0, 3] else 2
    add_writing_lines(doc, count=lines)

add_pull_quote(doc, "Competence does not protect you from complaints. Behaviour does.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 13 — MODULE 6: SCENARIO ASSESSMENT
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 6, "Application & Integration", "3:45 – 4:45 PM")

add_activity_box(doc, "✏️  SCENARIO ASSESSMENT — Individual Written",
                 bg_color=CREAM,
                 instructions="Write specifically — the exact words you would use, the exact steps you would take. 10 minutes. This is your personal record of where you are at the end of today.")
doc.add_paragraph()

assessment_scenarios = [
    (
        "SCENARIO 1 — Morning Greeting",
        "It's 9:15 AM. You're doing shift start rounds. Lakshmi, a machine operator on your team for "
        "8 months, is at her station. You've never formally acknowledged her in the mornings beyond a nod. "
        "What do you do and say today?",
        4
    ),
    (
        "SCENARIO 2 — Bystander Moment",
        "A male team member tells you, laughing, that he made a comment to Deepa during lunch that "
        "made her uncomfortable — something about her taking extra time in the washroom. "
        "He doesn't think it's a big deal. What do you do?",
        4
    ),
    (
        "SCENARIO 3 — Private Feedback",
        "You need to give critical feedback to Rekha — her rework rate this week is 3x the team average. "
        "It's 4:30 PM, the floor is winding down, and her male colleagues are nearby. "
        "Walk through exactly how you handle this — location, timing, specific words.",
        5
    ),
]
for s_title, s_text, s_lines in assessment_scenarios:
    as_t = doc.add_table(rows=1, cols=1)
    set_table_full_width(as_t)
    remove_table_borders(as_t)
    asc = as_t.cell(0, 0)
    set_cell_bg(asc, "2E74B5")
    asp = asc.paragraphs[0]
    asp.paragraph_format.space_before = Pt(4)
    asp.paragraph_format.space_after = Pt(4)
    asp.paragraph_format.left_indent = Pt(6)
    asr = asp.add_run(s_title)
    asr.bold = True
    asr.font.size = Pt(9)
    asr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    st_t = doc.add_table(rows=1, cols=1)
    set_table_full_width(st_t)
    remove_table_borders(st_t)
    stc = st_t.cell(0, 0)
    set_cell_bg(stc, GREY_FILL)
    stp = stc.paragraphs[0]
    stp.paragraph_format.left_indent = Pt(8)
    stp.paragraph_format.space_before = Pt(6)
    stp.paragraph_format.space_after = Pt(6)
    str_ = stp.add_run(s_text)
    str_.font.size = Pt(9)
    str_.font.color.rgb = BLACK

    add_writing_lines(doc, count=s_lines, label="My response:")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PAGE 14 — MODULE 7: POST-ASSESSMENT + PERSONAL ACTION PLAN
# ═══════════════════════════════════════════════════════════════════

add_module_banner(doc, 7, "Commitment & Action Planning", "4:45 – 5:00 PM")

add_activity_box(doc, "✏️  POST-ASSESSMENT — Revisit Your Baseline (Page 3)",
                 bg_color=CREAM,
                 instructions="Flip back to Page 3. Rate yourself again on the same five statements. Notice what changed — and what didn't.")
doc.add_paragraph()

post_t = doc.add_table(rows=2, cols=5)
set_table_full_width(post_t)
remove_table_borders(post_t)
for ci, header in enumerate(["Statement 1", "Statement 2", "Statement 3", "Statement 4", "Statement 5"]):
    cell = post_t.cell(0, ci)
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"S{ci+1}")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for ci in range(5):
    cell = post_t.cell(1, ci)
    set_cell_bg(cell, WHITE)
    set_cell_borders(cell,
                     top={'val': 'single', 'sz': 2, 'color': 'CCCCCC'},
                     bottom={'val': 'single', 'sz': 6, 'color': '2E74B5'},
                     left={'val': 'single', 'sz': 2, 'color': 'CCCCCC'},
                     right={'val': 'single', 'sz': 2, 'color': 'CCCCCC'})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("  / 5").font.size = Pt(8)

add_writing_lines(doc, count=2, label="What changed? Where are the gaps that remain? Those gaps anchor your commitments below.")

add_divider(doc)

add_section_header(doc, "My Personal Action Plan — 30 Days", color=NAVY, size=12)
add_body(doc,
    "Three commitments. Specific. Observable. Behavioural. Not 'I will be more respectful' — that's invisible.\n"
    "'I will greet every team member by name at shift start for the next 30 working days' — that's a commitment.",
    size=9, space_before=2, space_after=10)

commitment_colors = ["1F3864", "2E74B5", "D6E4F7"]
commitment_labels = [
    ("COMMITMENT 1 — DAILY PRACTICE",
     "Starting Monday, I will ____________________________________________\n"
     "every day for 30 working days."),
    ("COMMITMENT 2 — A BEHAVIOUR I WILL STOP",
     "I will no longer ________________________________________________\n"
     "when ____________________________________________________________."),
    ("COMMITMENT 3 — ONE PERSON'S EXPERIENCE I WILL IMPROVE",
     "For [picture them — no name needed], I will specifically\n"
     "________________________________________________________________."),
]

for ci_idx, (label, template) in enumerate(commitment_labels):
    c_table = doc.add_table(rows=1, cols=1)
    set_table_full_width(c_table)
    remove_table_borders(c_table)
    cc = c_table.cell(0, 0)
    set_cell_bg(cc, commitment_colors[ci_idx])
    cp = cc.paragraphs[0]
    cp.paragraph_format.space_before = Pt(5)
    cp.paragraph_format.space_after = Pt(5)
    cp.paragraph_format.left_indent = Pt(8)
    cr = cp.add_run(label)
    cr.bold = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if ci_idx < 2 else NAVY
    add_writing_lines(doc, count=4)

doc.add_paragraph()

# Accountability section
acc_table2 = doc.add_table(rows=1, cols=2)
set_table_full_width(acc_table2)
remove_table_borders(acc_table2)
for ci, (label, prompt) in enumerate([
    ("The person whose experience I'm committing to improve:", "(Picture them. No name needed.)"),
    ("30-Day Review — I will bring:", "One behaviour I changed. One example — good or hard."),
]):
    cell = acc_table2.cell(0, ci)
    set_cell_bg(cell, GREY_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(6)
    r1 = p.add_run(label + "\n")
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(prompt)
    r2.font.size = Pt(8)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_t = doc.add_table(rows=1, cols=2)
set_table_full_width(date_t)
remove_table_borders(date_t)
for ci, label in enumerate(["Signed:", "30-Day Review Date:"]):
    cell = date_t.cell(0, ci)
    set_cell_borders(cell, bottom={'val': 'single', 'sz': 6, 'color': '2E74B5'})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY

add_pull_quote(doc, "This is not a certificate. This is a choice.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# BACK COVER — DO'S & DON'TS QUICK REFERENCE CARD
# ═══════════════════════════════════════════════════════════════════

# Top bar
bc_bar = doc.add_table(rows=1, cols=1)
set_table_full_width(bc_bar)
remove_table_borders(bc_bar)
bc_cell = bc_bar.cell(0, 0)
set_cell_bg(bc_cell, "1F3864")
bcp = bc_cell.paragraphs[0]
bcp.paragraph_format.space_before = Pt(10)
bcp.paragraph_format.space_after = Pt(10)
bcp.paragraph_format.left_indent = Pt(8)
bcp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bcr = bcp.add_run("QUICK REFERENCE CARD — TAKE THIS TO YOUR FLOOR")
bcr.bold = True
bcr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
bcr.font.size = Pt(12)

doc.add_paragraph()

# Do's & Don'ts table
dd_table = doc.add_table(rows=8, cols=2)
set_table_full_width(dd_table)
remove_table_borders(dd_table)

for ci, (header, bg) in enumerate([("✔  DO", "27AE60"), ("✘  DON'T", "C0392B")]):
    cell = dd_table.cell(0, ci)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(header)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

dos_donts = [
    ("Greet consistently — same tone and name for everyone",        "Ignore or look past women team members in the morning"),
    ("Offer a seat for conversations longer than 2 minutes",         "Make her ask twice or justify a basic need"),
    ("Give all feedback in private — never on the open floor",       "Correct publicly in front of colleagues or the team"),
    ("Use SBI: Situation · Behaviour · Impact",                      "Use 'always' / 'never' / character attacks"),
    ("Listen whole — reflect before you respond",                    "Problem-solve over her before she finishes speaking"),
    ("Notice discomfort and ask simply: 'Are you okay?'",            "Wait to be told before acknowledging visible distress"),
    ("Address what you allow — silence = permission",                "Laugh off or ignore comments made by male colleagues"),
]
for ri, (do, dont) in enumerate(dos_donts):
    bg = "F0FFF4" if ri % 2 == 0 else WHITE
    dont_bg = "FFF0F0" if ri % 2 == 0 else WHITE
    for ci, (text, cell_bg) in enumerate([(do, bg), (dont, dont_bg)]):
        cell = dd_table.cell(ri + 1, ci)
        set_cell_bg(cell, cell_bg)
        set_cell_borders(cell, bottom={'val': 'single', 'sz': 2, 'color': 'DDDDDD'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = BLACK

doc.add_paragraph()

add_section_header(doc, "My Three Commitments — At a Glance", color=NAVY, size=11)
add_body(doc, "Write these in your own words for daily reference:", size=9, space_after=4)
for i in range(3):
    cm_t = doc.add_table(rows=1, cols=2)
    set_table_full_width(cm_t)
    remove_table_borders(cm_t)
    nc = cm_t.cell(0, 0)
    vc = cm_t.cell(0, 1)
    set_cell_bg(nc, "1F3864")
    set_cell_bg(vc, WHITE)
    set_cell_borders(vc, bottom={'val': 'single', 'sz': 6, 'color': '2E74B5'})
    np = nc.paragraphs[0]
    np.paragraph_format.space_before = Pt(12)
    np.paragraph_format.space_after = Pt(12)
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(str(i + 1))
    nr.bold = True
    nr.font.size = Pt(14)
    nr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_paragraph()
footer_t = doc.add_table(rows=1, cols=1)
set_table_full_width(footer_t)
remove_table_borders(footer_t)
fc = footer_t.cell(0, 0)
set_cell_bg(fc, "1F3864")
fp = fc.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(8)
fp.paragraph_format.space_after = Pt(8)
fr = fp.add_run("Preparatory Training for Managers  ·  Building Respectful Leadership in Manufacturing Environments")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ─────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────
output_path = "/Volumes/Dev/Course_Generator/courses/inclusive-management-shop-floor/participant-workbook.docx"
doc.save(output_path)
print(f"Workbook saved to: {output_path}")
